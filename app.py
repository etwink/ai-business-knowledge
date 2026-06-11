"""Main Streamlit application for Document Analysis System."""

import streamlit as st
import dataclasses
from pathlib import Path
from datetime import datetime
from collections import defaultdict

# Inject truststore into SSL context for better certificate handling (especially in enterprise environments)
import truststore
truststore.inject_into_ssl()

from src.analyzers import (
    ProcessDocumentBuilder,
    GapAnalyzer,
    AnalysisResult,
    ProcessDocument,
    GapAnalysis
)
from src.storage import AnalysisStorage
from src.pipeline import ConversationalAgent, DocumentUpdate, ProcessContextAgent
from src.pipeline.context_agent import AUDIENCE_LABELS, AUDIENCE_NOTES, AUDIENCE_INTERVIEW_NOTES, AUDIENCE_DESCRIPTIONS, AUDIENCE_GROUPS, RESPONSE_MODE_LABELS, DETAIL_LEVEL_LABELS
from src.rag import KnowledgeBase, RAGAgent
from src.llm_integration import llm_usage_tracker
import config


def _fmt_tokens(n: int) -> str:
    """Format a token count with 4 significant figures and K/M/B suffix above 100 K."""
    if n < 100_000:
        return f"{n:,}"
    if n < 1_000_000:
        return f"{n / 1_000:.4g}K"
    if n < 1_000_000_000:
        return f"{n / 1_000_000:.4g}M"
    return f"{n / 1_000_000_000:.4g}B"


def initialize_session():
    """Initialize session state variables."""
    if 'uploaded_files' not in st.session_state:
        st.session_state.uploaded_files = []
    if 'analyses' not in st.session_state:
        st.session_state.analyses = []
    if 'process_document' not in st.session_state:
        st.session_state.process_document = None
    if 'gap_analysis' not in st.session_state:
        st.session_state.gap_analysis = None
    if 'current_session' not in st.session_state:
        st.session_state.current_session = None
    if 'storage' not in st.session_state:
        st.session_state.storage = AnalysisStorage()
    # Bulk mode state
    if 'bulk_scanned' not in st.session_state:
        st.session_state.bulk_scanned = None        # ScannedDocuments
    if 'bulk_clusters' not in st.session_state:
        st.session_state.bulk_clusters = None       # list[DocumentCluster]
    if 'bulk_cluster_summaries' not in st.session_state:
        st.session_state.bulk_cluster_summaries = None  # list[ClusterSummary]
    if 'bulk_missing_deps' not in st.session_state:
        st.session_state.bulk_missing_deps = {}         # {missing_stem: [referencing_stems]}
    if 'bulk_dep_graph' not in st.session_state:
        st.session_state.bulk_dep_graph = {}            # {stem: [dep_dicts]}
    if 'bulk_dep_analyses' not in st.session_state:
        st.session_state.bulk_dep_analyses = []         # list[FileAnalysis]
    # Process context state (set before analysis to guide the LLM)
    if 'process_context' not in st.session_state:
        st.session_state.process_context = None    # ProcessContext | None
    # Chat agent state
    if 'chat_agent' not in st.session_state:
        st.session_state.chat_agent = None
    if 'chat_messages' not in st.session_state:
        st.session_state.chat_messages = []
    if 'document_updates' not in st.session_state:
        st.session_state.document_updates = {}  # doc_name -> list[DocumentUpdate]
    # Knowledge base / RAG chat state
    if 'rag_agent' not in st.session_state:
        st.session_state.rag_agent = None
    if 'rag_messages' not in st.session_state:
        st.session_state.rag_messages = []
    # Audience setting — controls writing style across all LLM outputs
    if 'doc_audience' not in st.session_state:
        st.session_state.doc_audience = "new_employee"
    if 'custom_audience_note' not in st.session_state:
        st.session_state.custom_audience_note = ""
    if '_sme_role_description' not in st.session_state:
        st.session_state._sme_role_description = ""
    # Response mode — controls format and depth of chat answers
    if 'response_mode' not in st.session_state:
        st.session_state.response_mode = "standard"
    # Detail level — controls thoroughness within a response
    if 'detail_level' not in st.session_state:
        st.session_state.detail_level = "standard"
    # Staging keys — hold unsaved selections until the user clicks Save Audience
    if '_audience_staging' not in st.session_state:
        st.session_state._audience_staging = st.session_state.doc_audience
    if '_custom_audience_staging' not in st.session_state:
        st.session_state._custom_audience_staging = st.session_state.custom_audience_note


def _save_session_settings() -> None:
    """Persist audience selection and process context for the current session."""
    session_name = st.session_state.get("current_session") or _ensure_session()
    ctx = st.session_state.get("process_context")
    settings = {
        "doc_audience": st.session_state.get("doc_audience", "new_employee"),
        "custom_audience_note": st.session_state.get("custom_audience_note", ""),
        "sme_role_description": st.session_state.get("_sme_role_description", ""),
        "process_context": {
            "foundation_document": ctx.foundation_document if ctx else None,
            "process_description": ctx.process_description if ctx else "",
            "additional_notes": ctx.additional_notes if ctx else "",
        } if ctx else None,
    }
    st.session_state.storage.save_settings(session_name, settings)


def _ensure_session() -> str:
    """Return current session name, auto-creating a timestamped one if not set."""
    if not st.session_state.current_session:
        st.session_state.current_session = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    llm_usage_tracker.start_session(st.session_state.current_session)
    return st.session_state.current_session


def _get_context_block() -> str:
    """Return the full context block (audience note + process context) for LLM injection."""
    audience = st.session_state.get("doc_audience", "new_employee")
    if audience == "custom":
        custom_text = st.session_state.get("custom_audience_note", "").strip()
        audience_note = (
            f"IMPORTANT — Audience: {custom_text}"
            if custom_text else ""
        )
    else:
        audience_note = AUDIENCE_NOTES.get(audience, "")
    ctx_block = (
        st.session_state.process_context.to_prompt_block()
        if st.session_state.process_context and st.session_state.process_context.is_set()
        else ""
    )
    if audience_note and ctx_block:
        return audience_note + "\n\n" + ctx_block
    return audience_note or ctx_block


def render_sidebar():
    """Render sidebar navigation and settings."""
    with st.sidebar:
        st.title("📚 Document Analysis")

        st.subheader("Session Management")
        
        # Session creation/loading
        col1, col2 = st.columns([2, 1])
        with col1:
            session_name = st.text_input(
                "Session name",
                value=st.session_state.current_session or "analysis_session",
                placeholder="e.g., payroll_system"
            )
        with col2:
            if st.button("📁 New", help="Create or load session"):
                st.session_state.current_session = session_name
                st.rerun()
        
        if st.session_state.current_session:
            st.success(f"Session: **{st.session_state.current_session}**")
        else:
            st.caption("A session will be created automatically on the first LLM action.")
        
        # Load previous sessions
        existing_sessions = st.session_state.storage.list_sessions()
        if existing_sessions:
            st.subheader("Previous Sessions")
            selected_session = st.selectbox(
                "Load session",
                existing_sessions,
                key="session_selector"
            )
            if st.button("📂 Load Session"):
                try:
                    session_data = st.session_state.storage.load_session(selected_session)
                    st.session_state.current_session = selected_session
                    
                    # Populate session data
                    if 'analyses' in session_data:
                        # Reconstruct AnalysisResult objects
                        st.session_state.analyses = [
                            AnalysisResult(**analysis) for analysis in session_data['analyses']
                        ]
                    
                    if 'process_document' in session_data:
                        doc = session_data['process_document']
                        st.session_state.process_document = ProcessDocument(
                            overview=doc.get('overview', ''),
                            integrated_processes=doc.get('integrated_processes', ''),
                            dependencies=doc.get('dependencies', ''),
                            data_flow=doc.get('data_flow', ''),
                            decision_points=doc.get('decision_points', ''),
                            systems_and_components=doc.get('systems_and_components', ''),
                            appendix=doc.get('appendix', ''),
                            process_flow_diagram=doc.get('process_flow_diagram', ''),
                            process_flow_ascii=doc.get('process_flow_ascii', ''),
                            citations=doc.get('citations', {}),
                        )
                    
                    if 'gap_analysis' in session_data:
                        gaps = session_data['gap_analysis']
                        if gaps.get('format') == 'v2':
                            st.session_state.gap_analysis = GapAnalysis(
                                gaps_by_category=gaps.get('gaps_by_category', {}),
                                edge_cases=gaps.get('edge_cases', []),
                                resource_gaps=gaps.get('resource_gaps', []),
                                ranked_gaps=gaps.get('ranked_gaps', []),
                            )
                        else:
                            # Migrate old format: map fixed fields → gaps_by_category
                            old_cats = {}
                            _old_map = [
                                ('missing_steps', 'Missing Steps'),
                                ('undefined_dependencies', 'Undefined Dependencies'),
                                ('incomplete_transformations', 'Incomplete Transformations'),
                                ('missing_integrations', 'Missing Integrations'),
                                ('error_handling_gaps', 'Error Handling Gaps'),
                                ('security_gaps', 'Security Gaps'),
                            ]
                            for field_name, cat_label in _old_map:
                                vals = [v for v in gaps.get(field_name, [])
                                        if v and not v.startswith("No ")]
                                if vals:
                                    old_cats[cat_label] = vals
                            st.session_state.gap_analysis = GapAnalysis(
                                gaps_by_category=old_cats,
                                edge_cases=[],
                                resource_gaps=[v for v in gaps.get('resource_gaps', [])
                                              if v and not v.startswith("No ")],
                                ranked_gaps=[],
                            )
                    
                    if 'chat_messages' in session_data:
                        st.session_state.chat_messages = session_data['chat_messages']
                        st.session_state.chat_agent = None  # agent must be re-initialized

                    if 'document_updates' in session_data:
                        doc_updates = {}
                        for doc_name, updates in session_data['document_updates'].items():
                            doc_updates[doc_name] = [
                                DocumentUpdate(
                                    document_name=u['document_name'],
                                    version=u['version'],
                                    content=u['content'],
                                    gap_addressed=u.get('gap_addressed', ''),
                                )
                                for u in updates
                            ]
                        st.session_state.document_updates = doc_updates

                    if 'settings' in session_data:
                        _s = session_data['settings']
                        if 'doc_audience' in _s:
                            st.session_state.doc_audience = _s['doc_audience']
                            st.session_state._audience_staging = _s['doc_audience']
                        if 'custom_audience_note' in _s:
                            st.session_state.custom_audience_note = _s['custom_audience_note']
                            st.session_state._custom_audience_staging = _s['custom_audience_note']
                        if 'sme_role_description' in _s:
                            st.session_state._sme_role_description = _s['sme_role_description']
                        if _s.get('process_context'):
                            from src.pipeline.context_agent import ProcessContext
                            _pc = _s['process_context']
                            st.session_state.process_context = ProcessContext(
                                foundation_document=_pc.get('foundation_document'),
                                process_description=_pc.get('process_description', ''),
                                additional_notes=_pc.get('additional_notes', ''),
                            )

                    # Restore clusters and cluster summaries (expensive to rebuild)
                    _loaded_clusters = st.session_state.storage.load_clusters(selected_session)
                    if _loaded_clusters:
                        st.session_state.bulk_clusters = _loaded_clusters
                    _loaded_summaries = st.session_state.storage.load_cluster_summaries(selected_session)
                    if _loaded_summaries:
                        st.session_state.bulk_cluster_summaries = _loaded_summaries

                    llm_usage_tracker.load_session(selected_session)
                    st.success(f"Loaded session: {selected_session}")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error loading session: {str(e)}")

        st.divider()
        st.subheader("Navigation")
        # Show context and audience status in sidebar
        _aud_key = st.session_state.get("doc_audience", "new_employee")
        if _aud_key == "custom":
            _custom = st.session_state.get("custom_audience_note", "").strip()
            audience_label = _custom[:60] + ("…" if len(_custom) > 60 else "") if _custom else "Custom (not yet described)"
        else:
            audience_label = AUDIENCE_LABELS.get(_aud_key, "—")
        if st.session_state.process_context and st.session_state.process_context.is_set():
            ctx = st.session_state.process_context
            ctx_lines = []
            if ctx.foundation_document:
                ctx_lines.append(f"📄 Foundation: **{ctx.foundation_document}**")
            if ctx.process_description:
                preview = ctx.process_description[:80]
                ctx_lines.append(f"💬 {preview}{'…' if len(ctx.process_description) > 80 else ''}")
            ctx_lines.append(f"👥 Audience: **{audience_label}**")
            _aud_desc = AUDIENCE_DESCRIPTIONS.get(_aud_key, "")
            if _aud_desc:
                ctx_lines.append(f"*{_aud_desc}*")
            st.info("\n\n".join(ctx_lines))
        else:
            st.caption("No process context set — go to **Group Documents** to add one.")
            st.caption(f"👥 Audience: **{audience_label}**")
            _aud_desc = AUDIENCE_DESCRIPTIONS.get(_aud_key, "")
            if _aud_desc:
                st.caption(f"*{_aud_desc}*")

        page = st.radio(
            "Select a step:",
            [
                "Group Documents",
                "Dependency Graph",
                "Review Process Document",
                "Gap Analysis",
                "Gap-Filling Chat",
                "Knowledge Chat",
            ]
        )

        st.subheader("Settings")
        api_status = "✅ Configured" if config.AZURE_OPENAI_API_KEY else "❌ Not configured"
        st.write(f"Azure OpenAI: {api_status}")

        # LLM cost tracker
        st.divider()
        usage = llm_usage_tracker.get_totals()
        if usage["total_calls"] > 0:
            st.caption("**LLM Usage (this session)**")
            c1, c2 = st.columns(2)
            c1.metric("Input tokens", _fmt_tokens(usage['total_input_tokens']))
            c2.metric("Output tokens", _fmt_tokens(usage['total_output_tokens']))
            st.caption(
                f"~**${usage['approximate_cost_usd']:.4f}** "
                f"({usage['total_calls']} calls) "
                f"— saved to `logs/usage_{usage['session_name']}.json`"
            )
        else:
            st.caption("LLM usage: no calls yet this session.")

        return page


def render_group_documents_page():
    """Unified page: load files, define context, build clusters, run hierarchical analysis."""
    from src.pipeline import FolderScanner, ClusterBuilder, HierarchicalSummarizer

    st.header("🗂️ Group Documents")
    st.write(
        "Load your documents, describe the process they relate to, then group them into "
        "logical clusters before running analysis."
    )

    # ── Step 1: Load files ────────────────────────────────────────────────────
    st.subheader("Step 1 — Load Files")

    load_mode = st.radio(
        "How would you like to load files?",
        ["Scan a folder path", "Upload individual files"],
        horizontal=True,
        key="group_load_mode",
    )

    if load_mode == "Scan a folder path":
        default_paths = ", ".join(str(p) for p in config.DOCUMENTS_PATHS)
        raw_paths = st.text_input(
            "Document folder path(s) — comma-separated for multiple folders",
            value=default_paths,
            placeholder="e.g. C:/projects/cobol_source, C:/projects/business_docs",
            help="Set DOCUMENTS_PATH in your .env file to pre-fill this field.",
        )
        input_paths = [Path(p.strip()) for p in raw_paths.split(",") if p.strip()]
        missing = [str(p) for p in input_paths if not p.exists()]
        if missing:
            st.warning(f"Path(s) not found: {', '.join(missing)}")

        recursive = st.checkbox("Scan sub-folders recursively", value=True)

        if st.button("🔍 Scan Folder(s)", disabled=not input_paths or bool(missing)):
            scanner = FolderScanner()
            with st.spinner("Scanning…"):
                st.session_state.bulk_scanned = scanner.scan(input_paths, recursive=recursive)
                st.session_state.bulk_clusters = None
                st.session_state.bulk_cluster_summaries = None
                st.session_state.uploaded_files = []

    else:  # Upload individual files
        uploaded = st.file_uploader(
            "Choose files",
            accept_multiple_files=True,
            type=["txt", "py", "cob", "cbl", "cic", "cpy", "mps", "src", "ct1", "jcv", "prv", "docx", "doc", "xlsx", "xlsm", "xlsb", "html"],
        )
        if uploaded:
            for uf in uploaded:
                if uf.name not in [getattr(f, "name", str(f)) for f in st.session_state.uploaded_files]:
                    st.session_state.uploaded_files.append(uf)

        if st.session_state.uploaded_files:
            st.write(f"**{len(st.session_state.uploaded_files)} file(s) loaded:**")
            for idx, file in enumerate(st.session_state.uploaded_files):
                col_name, col_rm = st.columns([5, 1])
                col_name.caption(f"📄 {getattr(file, 'name', str(file))}")
                if col_rm.button("🗑️", key=f"rm_upload_{idx}"):
                    st.session_state.uploaded_files.pop(idx)
                    st.rerun()
        else:
            st.info("No files uploaded yet.")

    # Show scan results summary
    scanned = st.session_state.bulk_scanned
    if scanned:
        st.success(f"Found {scanned.summary()}")
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("COBOL files", len(scanned.cobol))
        col2.metric("Source code", len(scanned.code))
        col3.metric("Word docs", len(scanned.word))
        col4.metric("Excel files", len(scanned.excel))
        with st.expander("Show all discovered files"):
            if scanned.cobol:
                st.write("**COBOL**")
                for p in scanned.cobol: st.caption(str(p))
            if scanned.code:
                st.write("**Source Code** (Python, SQL, JS, etc.)")
                for p in scanned.code: st.caption(str(p))
            if scanned.word:
                st.write("**Word**")
                for p in scanned.word: st.caption(str(p))
            if scanned.excel:
                st.write("**Excel**")
                for p in scanned.excel: st.caption(str(p))

    files_ready = (scanned and scanned.total_count > 0) or bool(st.session_state.uploaded_files)

    # ── Step 2: Define Context & Audience — always visible ───────────────────
    st.divider()
    st.subheader("Step 2 — Define Context (Optional but Recommended)")
    st.write(
        "Providing context helps the LLM understand the business domain and produce "
        "more accurate clusters and document sections. You can reference a specific "
        "document by name and the system will read it automatically."
    )
    ctx = st.session_state.process_context
    if ctx and ctx.is_set():
        with st.expander("✅ Context is set — click to view or change", expanded=False):
            if ctx.foundation_document:
                st.write(f"**Foundation document:** {ctx.foundation_document}")
            if ctx.process_description:
                st.write("**Process description:**")
                st.write(ctx.process_description)
            if ctx.additional_notes:
                st.write("**Additional notes:**")
                st.write(ctx.additional_notes)
            if st.button("✏️ Clear and re-set context", key="group_clear_ctx"):
                st.session_state.process_context = None
                _save_session_settings()
                st.rerun()
    else:
        with st.expander("Set context now", expanded=files_ready):
            quick_ctx = st.text_area(
                "Describe the process — optionally reference a document by name as a foundation",
                placeholder=(
                    "e.g. 'Payroll processing system for a mid-size manufacturer on IBM z/OS. "
                    "Use PAYROLL_OVERVIEW.docx as the foundation document.'"
                ),
                height=100,
                key="group_quick_ctx",
            )
            if not files_ready:
                st.caption("ℹ️ Load files in Step 1 to enable automatic reference-document detection.")
            if st.button("🔍 Extract Context", key="group_save_ctx") and quick_ctx.strip():
                # Build list of Path objects for document detection
                if scanned:
                    available_paths = scanned.cobol + scanned.code + scanned.word + scanned.excel
                else:
                    available_paths = [
                        f for f in st.session_state.uploaded_files if isinstance(f, Path)
                    ]

                ctx_agent = ProcessContextAgent(available_paths)

                with st.spinner("Pass 1 — Checking for reference document…"):
                    detected_doc = ctx_agent._identify_reference_doc(quick_ctx.strip())

                doc_content = ""
                if detected_doc:
                    with st.spinner(f"Reading {detected_doc}…"):
                        doc_content = ctx_agent._read_doc(detected_doc)

                with st.spinner("Extracting context…"):
                    ctx_result = ctx_agent._extract_context(
                        quick_ctx.strip(), doc_content, detected_doc
                    )

                st.session_state.process_context = ctx_result
                _save_session_settings()
                if detected_doc:
                    st.info(f"📄 Reference document detected and read: **{detected_doc}**")
                st.success("Context saved.")
                st.rerun()

    # Audience selector — grouped button grid, stages changes until Save is clicked
    st.write("**Document Audience**")
    st.caption(
        "Sets the audience for this session. "
        "The process document is written **for** this audience. "
        "Gap analysis finds gaps **this audience (as SME) can answer**. "
        "Both chat agents talk **to** this audience."
    )
    _current_staging = st.session_state.get("_audience_staging", "new_employee")
    _group_cols = st.columns(len(AUDIENCE_GROUPS))
    for _col, (_gname, _gkeys) in zip(_group_cols, AUDIENCE_GROUPS.items()):
        with _col:
            st.caption(f"**{_gname}**")
            for _key in _gkeys:
                if st.button(
                    AUDIENCE_LABELS[_key],
                    key=f"_aud_btn_{_key}",
                    type="primary" if _current_staging == _key else "secondary",
                    use_container_width=True,
                ):
                    st.session_state._audience_staging = _key
                    st.rerun()
    _staged_desc = AUDIENCE_DESCRIPTIONS.get(_current_staging, "")
    if _staged_desc:
        st.caption(f"*{_staged_desc}*")
    if st.session_state._audience_staging == "custom":
        st.text_area(
            "Describe your audience",
            placeholder=(
                "e.g. 'External auditors with accounting knowledge but no IT background' "
                "or 'Call-centre agents who need step-by-step instructions without jargon'"
            ),
            height=80,
            key="_custom_audience_staging",
        )
    st.text_area(
        "Describe this SME's specific role (optional)",
        placeholder=(
            "e.g. 'Data entry operator on the enrollment screens — no access to job scheduler, "
            "system logs, or program internals' or 'Front-line claims processor who handles "
            "standard submissions but escalates complex cases'"
        ),
        help="Narrows gap analysis to what this specific person would know from their daily work. Helps filter out gaps only a developer or admin could answer.",
        height=80,
        key="_sme_role_description_staging",
        value=st.session_state.get("_sme_role_description", ""),
    )
    if st.button("💾 Save Audience"):
        st.session_state.doc_audience = st.session_state._audience_staging
        if st.session_state._audience_staging == "custom":
            st.session_state.custom_audience_note = st.session_state.get("_custom_audience_staging", "")
        st.session_state._sme_role_description = st.session_state.get("_sme_role_description_staging", "")
        _save_session_settings()
        st.rerun()
    # Show unsaved-changes warning when staging differs from committed value
    _saved_aud = st.session_state.get("doc_audience", "new_employee")
    _staged_aud = st.session_state.get("_audience_staging", _saved_aud)
    _saved_custom = st.session_state.get("custom_audience_note", "")
    _staged_custom = st.session_state.get("_custom_audience_staging", "")
    _has_unsaved = _staged_aud != _saved_aud or (
        _staged_aud == "custom" and _staged_custom != _saved_custom
    )
    if _has_unsaved:
        _saved_label = AUDIENCE_LABELS.get(_saved_aud, "—")
        if _saved_aud == "custom":
            _saved_label = (
                (_saved_custom[:50] + "…" if len(_saved_custom) > 50 else _saved_custom)
                if _saved_custom else "Custom (not described)"
            )
        st.caption(f"⚠️ Unsaved changes — currently saved: **{_saved_label}**")

    # ── Step 3: Build clusters ────────────────────────────────────────────────
    if files_ready or st.session_state.bulk_clusters:
        st.divider()
        st.subheader("Step 3 — Build Dependency Clusters")
        st.write(
            "Files are grouped into logical subsystem clusters. COBOL files are clustered "
            "by their CALL/COPY dependency graph; all other files are grouped by subject matter "
            "using the LLM, informed by the context you defined above."
        )

        ctx_block = _get_context_block()

        if st.button("🧩 Build Clusters", disabled=st.session_state.bulk_clusters is not None):
            builder = ClusterBuilder()
            with st.spinner("Analysing dependencies and clustering documents…"):
                try:
                    if scanned:
                        cobol = scanned.cobol
                        word = scanned.word
                        excel = scanned.excel
                        code = scanned.code
                    else:
                        # Uploaded files — sort by extension into buckets
                        import tempfile, shutil
                        tmp_dir = Path(tempfile.mkdtemp())
                        cobol, word, excel, code = [], [], [], []
                        word_exts = {".doc", ".docx"}
                        excel_exts = {".xlsx", ".xlsm", ".xlsb", ".xls"}
                        cobol_exts = {".cob", ".cbl", ".cic", ".cpy", ".mps", ".src", ".ct1", ".jcv", ".prv", ".cobol"}
                        for uf in st.session_state.uploaded_files:
                            ext = Path(getattr(uf, "name", str(uf))).suffix.lower()
                            dest = tmp_dir / getattr(uf, "name", str(uf))
                            if hasattr(uf, "read"):
                                dest.write_bytes(uf.read())
                            else:
                                shutil.copy(str(uf), str(dest))
                            if ext in cobol_exts: cobol.append(dest)
                            elif ext in word_exts: word.append(dest)
                            elif ext in excel_exts: excel.append(dest)
                            else: code.append(dest)

                    st.session_state.bulk_clusters = builder.build_clusters(
                        cobol_files=cobol,
                        word_files=word,
                        excel_files=excel,
                        code_files=code,
                        context_block=ctx_block,
                    )
                    st.session_state.bulk_missing_deps = builder.missing_dependencies
                    st.session_state.bulk_dep_graph = builder.dep_graph
                    st.session_state.bulk_dep_analyses = builder.dep_analyses
                    st.session_state.bulk_cluster_summaries = None
                    try:
                        _session = _ensure_session()
                        st.session_state.storage.save_clusters(_session, st.session_state.bulk_clusters)
                        _save_session_settings()
                    except Exception:
                        pass
                except Exception as e:
                    st.error(f"Clustering failed: {e}")

        if st.session_state.bulk_clusters:
            clusters = st.session_state.bulk_clusters
            st.success(f"Built {len(clusters)} clusters")
            for cl in clusters:
                label = {
                    "cobol": "🖥️ Source Code",
                    "mixed": "🔗 Code + Docs",
                    "docs": "📄 Documents",
                }.get(cl.cluster_type, cl.cluster_type)
                with st.expander(f"{label}: {cl.cluster_name} ({cl.file_count} files)"):
                    if cl.entry_point:
                        st.write(f"**Entry point:** {cl.entry_point}")
                    if cl.cobol_files:
                        st.write(f"**Source files ({len(cl.cobol_files)}):** " +
                                 ", ".join(p.name for p in cl.cobol_files[:20]) +
                                 ("…" if len(cl.cobol_files) > 20 else ""))
                    if cl.doc_files:
                        st.write(f"**Docs ({len(cl.doc_files)}):** " +
                                 ", ".join(p.name for p in cl.doc_files[:10]) +
                                 ("…" if len(cl.doc_files) > 10 else ""))
                    if cl.shared_doc_files:
                        st.write(f"**🌐 Shared reference docs ({len(cl.shared_doc_files)}):** " +
                                 ", ".join(p.name for p in cl.shared_doc_files[:10]) +
                                 ("…" if len(cl.shared_doc_files) > 10 else ""))

            missing_deps = st.session_state.get("bulk_missing_deps", {})
            if missing_deps:
                with st.expander(f"⚠️ Missing dependencies ({len(missing_deps)})", expanded=False):
                    st.caption(
                        "These names were referenced in COBOL source files (via COPY, CALL, "
                        "EXEC SQL INCLUDE, etc.) but no matching file was found in the scanned "
                        "catalogue. They may be external programs, system copybooks, or files "
                        "that were not included in the scan."
                    )
                    for dep_name, refs in sorted(missing_deps.items()):
                        ref_str = ", ".join(sorted(set(refs)))
                        st.write(f"**{dep_name}** — referenced by: {ref_str}")

            if st.button("🔄 Re-cluster (clear and rebuild)"):
                st.session_state.bulk_clusters = None
                st.session_state.bulk_cluster_summaries = None
                st.session_state.bulk_missing_deps = {}
                st.rerun()

    # ── Step 4: Summarise clusters ────────────────────────────────────────────
    if st.session_state.bulk_clusters:
        st.divider()
        st.subheader("Step 4 — Summarize Clusters")
        clusters = st.session_state.bulk_clusters
        n = len(clusters)
        st.write(
            f"Each of the {n} clusters will be summarized separately, then synthesized "
            "into a single process document. "
            "This makes ~{} LLM calls — expect several minutes for large sets.".format(
                sum(len(cl.cobol_files) + len(cl.doc_files) + 1 for cl in clusters)
            )
        )

        if st.session_state.bulk_cluster_summaries is None:
            from src.pipeline.hierarchical_summarizer import _MAX_CLUSTER_WORKERS
            st.caption(
                f"Clusters are processed {_MAX_CLUSTER_WORKERS} at a time. "
                "Each cluster also parallelizes its per-file summaries internally."
            )
            if st.button("▶️ Run Hierarchical Analysis"):
                summarizer = HierarchicalSummarizer()
                progress_bar = st.progress(0, text="Starting…")
                active_slot = st.empty()   # shows clusters currently in progress
                done_slot = st.empty()     # shows the most recently completed cluster

                def on_progress(completed_name, done, total, active):
                    frac = done / total if total else 0
                    progress_bar.progress(frac, text=f"Completed {done}/{total} clusters")
                    done_slot.markdown(f"✅ **{completed_name}** — complete")
                    if active:
                        active_slot.markdown(
                            "⏳ In progress: " + " · ".join(f"**{n}**" for n in active)
                        )
                    else:
                        active_slot.empty()

                ctx_block = _get_context_block()
                try:
                    summaries = summarizer.summarize_all(
                        clusters,
                        progress_callback=on_progress,
                        context_block=ctx_block,
                    )
                    st.session_state.bulk_cluster_summaries = summaries
                    # ConversationalAgent and _map_gaps_to_clusters expect AnalysisResult
                    # objects; convert each ClusterSummary to a minimal AnalysisResult so
                    # both interfaces stay consistent without refactoring their signatures.
                    st.session_state.analyses = [
                        AnalysisResult(
                            document_name=cs.cluster_name,
                            summary=cs.summary,
                            key_processes=cs.key_processes,
                            systems_mentioned=cs.systems_mentioned,
                            technical_details=[f"{cs.file_count} files in cluster"],
                        )
                        for cs in summaries
                    ]
                    progress_bar.progress(1.0, text="Done!")
                    active_slot.empty()
                    done_slot.empty()
                    try:
                        _session = _ensure_session()
                        st.session_state.storage.save_analyses(_session, st.session_state.analyses)
                        st.session_state.storage.save_cluster_summaries(_session, summaries)
                        _save_session_settings()
                        st.success(f"Summarized {len(summaries)} clusters. Saved to session: **{st.session_state.current_session}**")
                    except Exception as e:
                        st.success(f"Summarized {len(summaries)} clusters.")
                        st.warning(f"Could not save analyses: {e}")
                except Exception as e:
                    st.error(f"Summarization failed: {e}")
        else:
            st.success(f"{len(st.session_state.bulk_cluster_summaries)} cluster summaries ready.")
            with st.expander("Preview cluster summaries"):
                for cs in st.session_state.bulk_cluster_summaries:
                    st.write(f"**{cs.cluster_name}** ({cs.file_count} files)")
                    st.write(cs.summary)
                    st.divider()

    # ── Step 5: Build process document ───────────────────────────────────────
    if st.session_state.bulk_cluster_summaries:
        st.divider()
        st.subheader("Step 5 — Build Process Document")
        st.write(
            "Synthesize the cluster summaries into one integrated process document. "
            "This is the same document you'll see in the Review Process Document page."
        )

        if st.button("📝 Build Process Document"):
            builder = ProcessDocumentBuilder()
            section_progress = st.progress(0, text="Starting…")
            section_status = st.empty()

            def on_section(label, idx, total):
                section_progress.progress(idx / total, text=f"Writing section {idx}/{total}")
                section_status.write(f"Generating: **{label}**")

            ctx_block = _get_context_block()
            try:
                process_doc = builder.build_from_cluster_summaries(
                    st.session_state.bulk_cluster_summaries,
                    progress_callback=on_section,
                    context_block=ctx_block,
                )
                section_progress.progress(1.0, text="Done!")
                section_status.empty()
                st.session_state.process_document = process_doc

                try:
                    _session = _ensure_session()
                    st.session_state.storage.save_process_document(_session, process_doc, version="v1")
                    _save_session_settings()
                    st.success(f"Process document built and saved to session: **{st.session_state.current_session}**. Navigate to **Review Process Document** to continue.")
                except Exception as save_err:
                    st.success("Process document built! Navigate to **Review Process Document** to continue.")
                    st.warning(f"Could not save: {save_err}")
                st.balloons()
            except Exception as e:
                st.error(f"Process document build failed: {e}")

        if st.session_state.process_document:
            st.info(
                "✅ Process document is ready. Use the sidebar to navigate to "
                "**Review Process Document → Gap Analysis → Questions → …**"
            )




# ---------------------------------------------------------------------------
# Dependency Graph page
# ---------------------------------------------------------------------------

_CLUSTER_TYPE_COLORS = {
    "cobol": "#4C72B0",
    "mixed": "#55A868",
    "docs":  "#DD8452",
}

_PLOTLY_PALETTE = [
    "#636EFA", "#EF553B", "#00CC96", "#AB63FA", "#FFA15A",
    "#19D3F3", "#FF6692", "#B6E880", "#FF97FF", "#FECB52",
    "#1F77B4", "#FF7F0E", "#2CA02C", "#D62728", "#9467BD",
    "#8C564B", "#E377C2", "#7F7F7F", "#BCBD22", "#17BECF",
]

_SIBLING_PALETTE = ["#E74C3C", "#9B59B6", "#F39C12", "#1ABC9C", "#3498DB", "#E67E22"]


def _hierarchical_positions(nodes: list, edges: list) -> dict:
    """Return {node: (x, y)} in a top-down tree layout.

    Roots (in-degree 0) sit at y=0; each dependency level descends by 1.5 units.
    Nodes at the same level are evenly spaced on the x-axis.
    Falls back to a circular layout when networkx is unavailable.
    """
    try:
        import networkx as nx
        from collections import defaultdict, deque

        node_set = set(nodes)
        G = nx.DiGraph()
        G.add_nodes_from(nodes)
        for s, t in edges:
            if s in node_set and t in node_set:
                G.add_edge(s, t)

        # Longest-path level from any root (BFS, keep updating to max)
        in_deg = dict(G.in_degree())
        level: dict = {}
        roots = [n for n in nodes if in_deg.get(n, 0) == 0]
        if not roots:
            roots = [nodes[0]] if nodes else []

        q = deque()
        for r in roots:
            level[r] = 0
            q.append(r)
        while q:
            node = q.popleft()
            for _, child in G.out_edges(node):
                new_lv = level[node] + 1
                if child not in level or level[child] < new_lv:
                    level[child] = new_lv
                    q.append(child)

        # Disconnected nodes go one row below the deepest connected node
        max_lv = max(level.values(), default=0)
        for n in nodes:
            if n not in level:
                level[n] = max_lv + 1

        # Group and sort nodes within each level for a stable, readable layout
        lv_to_nodes: dict = defaultdict(list)
        for n in nodes:
            lv_to_nodes[level[n]].append(n)

        max_width = max(len(ns) for ns in lv_to_nodes.values())
        positions = {}
        for lv, ns in lv_to_nodes.items():
            for i, n in enumerate(sorted(ns)):
                x = (i + 0.5) * max_width / len(ns) - max_width / 2
                positions[n] = (x, -lv * 1.5)
        return positions

    except ImportError:
        import math
        n = max(len(nodes), 1)
        return {
            node: (math.cos(2 * math.pi * i / n), math.sin(2 * math.pi * i / n))
            for i, node in enumerate(nodes)
        }


def _edge_traces(positions: dict, edges: list, color: str = "rgba(150,150,150,0.35)", width: float = 1.0):
    import plotly.graph_objects as go
    result = []
    for src, tgt in edges:
        if src not in positions or tgt not in positions:
            continue
        x0, y0 = positions[src]
        x1, y1 = positions[tgt]
        result.append(go.Scatter(
            x=[x0, x1, None], y=[y0, y1, None],
            mode="lines",
            line=dict(color=color, width=width),
            hoverinfo="none",
            showlegend=False,
        ))
    return result


def _render_cluster_graph(
    cluster_info: dict,
    inter_edges: dict,
    sibling_groups: dict | None = None,  # parent_cluster_id → [child_cluster_ids]
):
    import plotly.graph_objects as go

    sibling_groups = sibling_groups or {}

    # Assign one border color per sibling group
    sibling_border_color: dict[str, str] = {}
    for i, siblings in enumerate(sibling_groups.values()):
        color = _SIBLING_PALETTE[i % len(_SIBLING_PALETTE)]
        for cid in siblings:
            sibling_border_color[cid] = color

    nodes = list(cluster_info.keys())
    edge_list = list(inter_edges.keys())
    positions = _hierarchical_positions(nodes, edge_list)

    traces: list = []
    annotations: list = []

    # --- Sibling dotted lines (drawn first, so they sit beneath nodes) ---
    for parent_id, siblings in sibling_groups.items():
        group_color = sibling_border_color.get(siblings[0], "#888") if siblings else "#888"
        for i, a in enumerate(siblings):
            for b in siblings[i + 1:]:
                if a in positions and b in positions:
                    x0, y0 = positions[a]
                    x1, y1 = positions[b]
                    traces.append(go.Scatter(
                        x=[x0, x1, None], y=[y0, y1, None],
                        mode="lines",
                        line=dict(color=group_color, width=1.5, dash="dot"),
                        hoverinfo="none",
                        showlegend=False,
                    ))

    # --- Dependency edges ---
    max_w = max(inter_edges.values(), default=1)
    for (src, tgt), count in inter_edges.items():
        if src not in positions or tgt not in positions:
            continue
        x0, y0 = positions[src]
        x1, y1 = positions[tgt]
        traces.append(go.Scatter(
            x=[x0, x1, None], y=[y0, y1, None],
            mode="lines",
            line=dict(color="rgba(120,120,120,0.45)", width=1 + 3 * count / max_w),
            hoverinfo="none",
            showlegend=False,
        ))
        annotations.append(dict(
            ax=x0, ay=y0, x=x1, y=y1,
            xref="x", yref="y", axref="x", ayref="y",
            showarrow=True, arrowhead=2, arrowsize=1.2,
            arrowcolor="rgba(100,100,100,0.55)", arrowwidth=1.5,
        ))

    # --- Nodes ---
    node_x, node_y, labels, hovers = [], [], [], []
    fill_colors, border_colors, border_widths, sizes = [], [], [], []
    for cid, info in cluster_info.items():
        if cid not in positions:
            continue
        x, y = positions[cid]
        node_x.append(x)
        node_y.append(y)
        labels.append(info["name"])
        hovers.append(f"<b>{info['name']}</b><br>Type: {info['type']}<br>Files: {info['file_count']}")
        fill_colors.append(_CLUSTER_TYPE_COLORS.get(info["type"], "#888"))
        bc = sibling_border_color.get(cid)
        border_colors.append(bc if bc else "#ffffff")
        border_widths.append(3 if bc else 1.5)
        sizes.append(max(22, 18 + info["file_count"] * 2))

    traces.append(go.Scatter(
        x=node_x, y=node_y,
        mode="markers+text",
        text=labels,
        textposition="top center",
        textfont=dict(color="#1a1a1a", size=11),
        marker=dict(
            color=fill_colors,
            size=sizes,
            line=dict(color=border_colors, width=border_widths),
        ),
        hovertext=hovers,
        hoverinfo="text",
        showlegend=False,
    ))

    # --- Legend: cluster types ---
    type_labels = {"cobol": "Source Code", "mixed": "Code + Docs", "docs": "Documents"}
    for ct in sorted({info["type"] for info in cluster_info.values()}):
        traces.append(go.Scatter(
            x=[None], y=[None], mode="markers",
            marker=dict(color=_CLUSTER_TYPE_COLORS.get(ct, "#888"), size=12),
            name=type_labels.get(ct, ct),
        ))

    # --- Legend: sibling groups ---
    for parent_id, siblings in sibling_groups.items():
        group_color = sibling_border_color.get(siblings[0], "#888") if siblings else "#888"
        traces.append(go.Scatter(
            x=[None], y=[None], mode="markers",
            marker=dict(
                color="white", size=12, symbol="square",
                line=dict(color=group_color, width=3),
            ),
            name=f"Split: {parent_id}",
        ))

    return go.Figure(
        data=traces,
        layout=go.Layout(
            height=540,
            margin=dict(l=20, r=20, t=40, b=20),
            xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
            yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
            plot_bgcolor="white", paper_bgcolor="white",
            annotations=annotations,
            legend=dict(orientation="h", yanchor="bottom", y=1.01, xanchor="right", x=1,
                        font=dict(color="#1a1a1a")),
            hovermode="closest",
        ),
    )


def _render_file_graph(
    nodes: list, edges: list, file_to_cluster: dict,
    cluster_info: dict, cluster_summaries: list,
):
    import plotly.graph_objects as go

    if not nodes:
        return go.Figure()

    positions = _hierarchical_positions(nodes, edges)
    traces = _edge_traces(positions, edges)

    color_map = {
        cs.cluster_id: _PLOTLY_PALETTE[i % len(_PLOTLY_PALETTE)]
        for i, cs in enumerate(cluster_summaries)
    }

    grouped: dict[str, list] = {}
    unresolved = []
    for node in nodes:
        cid = file_to_cluster.get(node)
        if cid:
            grouped.setdefault(cid, []).append(node)
        else:
            unresolved.append(node)

    for cid, members in grouped.items():
        info = cluster_info.get(cid, {})
        color = color_map.get(cid, "#888")
        xs = [positions[n][0] for n in members if n in positions]
        ys = [positions[n][1] for n in members if n in positions]
        traces.append(go.Scatter(
            x=xs, y=ys, mode="markers",
            marker=dict(color=color, size=10, line=dict(color="white", width=1)),
            hovertext=[f"<b>{n}</b><br>Cluster: {info.get('name', cid)}" for n in members if n in positions],
            hoverinfo="text",
            name=info.get("name", cid),
        ))

    if unresolved:
        xs = [positions[n][0] for n in unresolved if n in positions]
        ys = [positions[n][1] for n in unresolved if n in positions]
        traces.append(go.Scatter(
            x=xs, y=ys, mode="markers",
            marker=dict(color="#cccccc", size=8, symbol="diamond"),
            hovertext=[f"<b>{n}</b><br>Referenced but not in scanned folder" for n in unresolved if n in positions],
            hoverinfo="text",
            name="Missing / Unresolved",
        ))

    return go.Figure(
        data=traces,
        layout=go.Layout(
            height=640,
            margin=dict(l=20, r=20, t=20, b=20),
            xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
            yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
            plot_bgcolor="white", paper_bgcolor="white",
            hovermode="closest",
            legend=dict(x=1.01, y=1, font=dict(size=10), bgcolor="rgba(255,255,255,0.8)"),
        ),
    )


def render_dependency_graph_page():
    from collections import defaultdict
    st.header("🔗 Dependency Graph")

    cluster_summaries = st.session_state.get("bulk_cluster_summaries") or []
    dep_graph = st.session_state.get("bulk_dep_graph") or {}

    if not cluster_summaries and not dep_graph:
        st.info(
            "No dependency data yet. Complete **Group Documents → Step 3 (Build Clusters)** "
            "and **Step 4 (Summarize Clusters)** first."
        )
        return

    try:
        import plotly.graph_objects as go  # noqa: F401
    except ImportError:
        st.error(
            "`plotly` is not installed. Run `pip install plotly networkx` in your environment "
            "and restart the app."
        )
        return

    # file stem (upper-case) → cluster_id
    file_to_cluster: dict[str, str] = {}
    for cs in cluster_summaries:
        for f in cs.source_files:
            file_to_cluster[Path(f).stem.upper()] = cs.cluster_id

    cluster_info: dict[str, dict] = {
        cs.cluster_id: {"name": cs.cluster_name, "file_count": cs.file_count, "type": cs.cluster_type}
        for cs in cluster_summaries
    }

    inter_cluster_edges: dict[tuple, int] = defaultdict(int)
    file_edges: list[tuple] = []
    for src_stem, deps in dep_graph.items():
        src_cid = file_to_cluster.get(src_stem)
        for dep in deps:
            tgt = dep["target"]
            file_edges.append((src_stem, tgt))
            tgt_cid = file_to_cluster.get(tgt)
            if src_cid and tgt_cid and src_cid != tgt_cid:
                inter_cluster_edges[(src_cid, tgt_cid)] += 1

    # ── Cluster graph ────────────────────────────────────────────────────────
    if cluster_summaries:
        # Build sibling groups: clusters that were split from the same parent
        sibling_groups: dict[str, list[str]] = defaultdict(list)
        for cs in cluster_summaries:
            if getattr(cs, "parent_cluster_id", None):
                sibling_groups[cs.parent_cluster_id].append(cs.cluster_id)

        st.subheader("Cluster Dependency Graph")
        st.caption(
            "Each node is one cluster — size reflects file count. "
            "Arrows point **dependent → dependency** (A → B: something in A calls/copies B). "
            "Blue = source code only · Green = code + docs · Orange = documents only. "
            "Colored borders + dotted lines show clusters split from the same parent."
        )
        st.plotly_chart(
            _render_cluster_graph(cluster_info, dict(inter_cluster_edges), sibling_groups=dict(sibling_groups)),
            use_container_width=True,
        )

    st.divider()

    # ── File graph ───────────────────────────────────────────────────────────
    st.subheader("File Dependency Graph")

    filter_options = ["All clusters"] + [cs.cluster_name for cs in cluster_summaries]
    selected = st.selectbox("Scope to cluster", filter_options, key="dep_graph_filter")

    if selected != "All clusters":
        target_cs = next((cs for cs in cluster_summaries if cs.cluster_name == selected), None)
        if target_cs:
            cluster_files = {Path(f).stem.upper() for f in target_cs.source_files}
            visible = set(cluster_files)
            for stem in cluster_files:
                for dep in dep_graph.get(stem, []):
                    visible.add(dep["target"])
            for src, deps in dep_graph.items():
                if any(dep["target"] in cluster_files for dep in deps):
                    visible.add(src)
            filtered_edges = [(s, t) for s, t in file_edges if s in visible or t in visible]
            display_nodes = list(visible)
            external = len(visible) - len(cluster_files)
            st.caption(
                f"**{len(cluster_files)} files** in *{selected}* "
                f"+ {external} external neighbor(s). "
                "Nodes outside the selected cluster show the cluster boundary."
            )
        else:
            display_nodes = list(file_to_cluster.keys())
            filtered_edges = file_edges
    else:
        all_referenced = {t for _, t in file_edges}
        display_nodes = list(set(file_to_cluster.keys()) | all_referenced)
        filtered_edges = file_edges
        n_missing = len(all_referenced - set(file_to_cluster.keys()))
        st.caption(
            f"All **{len(display_nodes)} files** across **{len(cluster_summaries)} clusters**. "
            + (f"**{n_missing}** gray diamond node(s) are referenced but not in the scanned folder." if n_missing else "")
        )

    if display_nodes:
        st.plotly_chart(
            _render_file_graph(display_nodes, filtered_edges, file_to_cluster, cluster_info, cluster_summaries),
            use_container_width=True,
        )
    else:
        st.info("No file dependency data to display for this selection.")

    # ── Missing dependencies summary ─────────────────────────────────────────
    missing_deps = st.session_state.get("bulk_missing_deps") or {}
    if missing_deps:
        with st.expander(f"⚠️ {len(missing_deps)} referenced file(s) not found in scanned folder"):
            st.caption("These stems appear in COPY/CALL/DSN statements but have no matching scanned file.")
            for stem, referencing in sorted(missing_deps.items()):
                st.write(f"**{stem}** ← {', '.join(sorted(referencing))}")


def render_process_document_page():
    """Render process document review page."""
    st.header("📖 Process Document")

    has_summaries = bool(st.session_state.bulk_cluster_summaries)
    has_doc = bool(st.session_state.process_document)

    if not has_summaries and not has_doc:
        st.warning(
            "No cluster summaries available. Complete **Group Documents** steps 3 and 4 first, "
            "or load a session that already has a process document."
        )
        return

    if has_summaries:
        if st.button("Generate Process Document", type="primary"):
            builder = ProcessDocumentBuilder()
            section_progress = st.progress(0, text="Starting…")
            section_status = st.empty()

            def on_pd_section(label, idx, total):
                section_progress.progress(idx / total, text=f"Writing section {idx}/{total}")
                section_status.write(f"Generating: **{label}**")

            ctx_block = _get_context_block()
            try:
                st.session_state.process_document = builder.build_from_cluster_summaries(
                    st.session_state.bulk_cluster_summaries,
                    progress_callback=on_pd_section,
                    context_block=ctx_block,
                )
                section_progress.progress(1.0, text="Done!")
                section_status.empty()

                try:
                    st.session_state.storage.save_process_document(
                        _ensure_session(),
                        st.session_state.process_document,
                        version="v1"
                    )
                    st.success(f"✅ Process document generated and saved to session: **{st.session_state.current_session}**")
                except Exception as e:
                    st.success("✅ Process document generated!")
                    st.warning(f"Could not save: {e}")
            except Exception as e:
                st.error(f"Error generating process document: {str(e)}")

    if st.session_state.process_document:
        doc = st.session_state.process_document

        st.subheader("Overview")
        st.write(doc.overview)

        st.divider()
        col1, col2 = st.columns(2)

        with col1:
            st.subheader("Integrated Processes")
            st.write(doc.integrated_processes)

            st.subheader("Data Flow")
            st.write(doc.data_flow)

        with col2:
            st.subheader("Dependencies")
            st.write(doc.dependencies)

            st.subheader("Decision Points")
            st.write(doc.decision_points)

        st.divider()
        st.subheader("Systems & Components")
        st.write(doc.systems_and_components)

        if doc.appendix:
            st.divider()
            st.subheader("Appendix")
            st.write(doc.appendix)

        if doc.process_flow_diagram:
            st.divider()
            st.subheader("📊 Process Flow Diagram")
            tab_interactive, tab_png = st.tabs(["Interactive (Mermaid)", "Static PNG"])
            with tab_interactive:
                _render_mermaid(doc.process_flow_diagram, height=560)
                with st.expander("View / copy Mermaid source"):
                    st.code(doc.process_flow_diagram, language="text")
                if st.button("🔧 Regenerate Diagram", help="If you see a syntax error above, click to ask the AI to repair the diagram"):
                    with st.spinner("Repairing diagram syntax…"):
                        fixed = ProcessDocumentBuilder()._repair_mermaid(doc.process_flow_diagram)
                    updated = dataclasses.replace(st.session_state.process_document, process_flow_diagram=fixed)
                    st.session_state.process_document = updated
                    try:
                        st.session_state.storage.save_process_document(
                            _ensure_session(), updated, version="v1"
                        )
                    except Exception:
                        pass
                    st.rerun()
            with tab_png:
                with st.spinner("Rendering PNG…"):
                    png_bytes = _mermaid_to_png(doc.process_flow_diagram)
                if png_bytes:
                    st.image(png_bytes, use_container_width=True)
                    st.download_button(
                        "📥 Download diagram PNG",
                        data=png_bytes,
                        file_name="process_flow.png",
                        mime="image/png",
                    )
                else:
                    st.info("PNG rendering unavailable — check network access to mermaid.ink. The Word export will include the Mermaid source as a fallback.")

        # Export buttons
        st.divider()
        appendix_section = f"\n\n## Appendix\n{doc.appendix}" if doc.appendix else ""
        diagram_section = (
            f"\n\n## Process Flow Diagram\n```mermaid\n{doc.process_flow_diagram}\n```"
            if doc.process_flow_diagram else ""
        )
        markdown = f"""# Process Document

## Overview
{doc.overview}

## Integrated Processes
{doc.integrated_processes}

## Dependencies
{doc.dependencies}

## Data Flow
{doc.data_flow}

## Decision Points
{doc.decision_points}

## Systems & Components
{doc.systems_and_components}{appendix_section}{diagram_section}
"""
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="📥 Download as Markdown",
                data=markdown,
                file_name="process_document.md",
                mime="text/markdown",
            )
        with col2:
            st.download_button(
                label="📄 Download as Word",
                data=_generate_word_doc(doc),
                file_name="process_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


def _map_gaps_to_clusters(gap_descriptions: list[str], analyses: list) -> dict[str, str]:
    """Map each gap description to the best-matching cluster/analysis name via word overlap.

    Returns a dict of {gap_description: cluster_name}. Falls back to the first analysis
    when there is no meaningful overlap (so the caller always gets a name).
    """
    import re
    if not analyses or not gap_descriptions:
        return {}
    result: dict[str, str] = {}
    stopwords = {"the", "a", "an", "is", "are", "was", "were", "or", "and", "of",
                 "in", "to", "for", "with", "not", "this", "that", "has", "have",
                 "it", "its", "be", "by", "at", "on", "as", "no", "gaps", "gap",
                 "missing", "undefined", "incomplete", "each", "identified"}
    for gap in gap_descriptions:
        gap_words = set(re.findall(r'\w+', gap.lower())) - stopwords
        best_name = analyses[0].document_name
        best_score = -1
        for a in analyses:
            candidate = (
                a.summary + " "
                + " ".join(a.key_processes)
                + " " + " ".join(a.systems_mentioned)
            ).lower()
            overlap = len(gap_words & (set(re.findall(r'\w+', candidate)) - stopwords))
            if overlap > best_score:
                best_score, best_name = overlap, a.document_name
        result[gap] = best_name
    return result


def render_gap_analysis_page():
    """Render gap analysis page."""
    st.header("⚠️ Gap Analysis")

    if not st.session_state.process_document:
        st.warning("Please generate process document first")
        return

    if st.button("Analyze Gaps", type="primary"):
        gap_analyzer = GapAnalyzer()
        _audience = st.session_state.get("doc_audience", "new_employee")
        _sme_role = st.session_state.get("_sme_role_description", "").strip()
        if _audience == "custom":
            _base = st.session_state.get("custom_audience_note", "")
            _audience_note = (_base + "\n\n" + _sme_role).strip() if _sme_role else _base
        else:
            _audience_note = _sme_role
        _cluster_edge_cases = []
        if st.session_state.bulk_cluster_summaries:
            for _cs in st.session_state.bulk_cluster_summaries:
                if hasattr(_cs, "edge_cases"):
                    _cluster_edge_cases.extend(_cs.edge_cases)

        with st.spinner("Analyzing gaps — identifying gaps…"):
            st.session_state.gap_analysis = gap_analyzer.analyze_gaps(
                st.session_state.process_document,
                context_block=_get_context_block(),
                audience_key=_audience,
                audience_note=_audience_note,
                cluster_edge_cases=_cluster_edge_cases or None,
            )

        _fix_count = 0
        if st.session_state.bulk_cluster_summaries:
            with st.spinner("Verifying gaps against source documents…"):
                _file_path_map: dict = {}
                _bulk_scanned = st.session_state.get("bulk_scanned")
                if _bulk_scanned:
                    for _p in (
                        _bulk_scanned.cobol
                        + _bulk_scanned.code
                        + _bulk_scanned.word
                        + _bulk_scanned.excel
                    ):
                        _file_path_map[_p.name] = _p
                for _f in st.session_state.get("uploaded_files", []):
                    if isinstance(_f, Path):
                        _file_path_map[_f.name] = _f

                (
                    st.session_state.gap_analysis,
                    st.session_state.process_document,
                    st.session_state.bulk_cluster_summaries,
                    _fix_count,
                ) = gap_analyzer.verify_and_fix_gaps(
                    st.session_state.gap_analysis,
                    st.session_state.process_document,
                    st.session_state.bulk_cluster_summaries,
                    _file_path_map,
                )

        try:
            st.session_state.storage.save_gap_analysis(
                _ensure_session(),
                st.session_state.gap_analysis,
                version="v1"
            )
            if _fix_count > 0:
                try:
                    st.session_state.storage.save_process_document(
                        _ensure_session(),
                        st.session_state.process_document,
                        version="v1",
                    )
                    st.session_state.storage.save_cluster_summaries(
                        _ensure_session(),
                        st.session_state.bulk_cluster_summaries,
                    )
                except Exception:
                    pass
                st.info(
                    f"During gap analysis, **{_fix_count} documentation miss(es)** were identified "
                    "and patched into cluster summaries and the process document. "
                    "We recommend **rebuilding your process document** to ensure full consistency."
                )
            st.success(f"✅ Gap analysis complete and saved to session: **{st.session_state.current_session}**")
        except Exception as e:
            st.error(f"Error saving gap analysis: {str(e)}")

    if st.session_state.gap_analysis:
        gaps = st.session_state.gap_analysis

        # Collect all gap descriptions for the cluster-map helper
        _all_descs = (
            [d for items in gaps.gaps_by_category.values() for d in items]
            + gaps.edge_cases
            + gaps.resource_gaps
        )
        gap_cluster_map = _map_gaps_to_clusters(_all_descs, st.session_state.analyses)

        def _gap_item(item: str) -> None:
            cluster = gap_cluster_map.get(item, "")
            badge = f" — *{cluster}*" if cluster else ""
            st.markdown(f"- {item}{badge}")

        def _importance_badge(score: int) -> str:
            if score >= 8:
                return f"🔴 **{score}/10** Critical"
            elif score >= 6:
                return f"🟠 **{score}/10** High"
            elif score >= 4:
                return f"🟡 **{score}/10** Medium"
            else:
                return f"🟢 **{score}/10** Low"

        tab_ranked, tab_categories, tab_edge, tab_resource = st.tabs([
            "🏆 Ranked by Importance",
            "📋 By Category",
            "⚡ Edge Cases",
            "🧑‍💼 Resource Gaps",
        ])

        with tab_ranked:
            if gaps.ranked_gaps:
                _true_gaps = [g for g in gaps.ranked_gaps if not g.get("is_doc_miss", False)]
                _doc_misses = [g for g in gaps.ranked_gaps if g.get("is_doc_miss", False)]
                st.caption(
                    f"{len(_true_gaps)} gaps ranked by business importance — "
                    "🔴 Critical · 🟠 High · 🟡 Medium · 🟢 Low"
                    + (f" · {len(_doc_misses)} documentation misses" if _doc_misses else "")
                )
                for g in _true_gaps:
                    score = g.get("importance", 5)
                    cat = g.get("category", "")
                    desc = g.get("description", "")
                    cluster = gap_cluster_map.get(desc, "")
                    badge = f" — *{cluster}*" if cluster else ""
                    st.markdown(
                        f"{_importance_badge(score)} &nbsp; `{cat}` — {desc}{badge}",
                        unsafe_allow_html=True,
                    )
                if _doc_misses:
                    with st.expander(f"📋 Documentation misses ({len(_doc_misses)}) — covered in source docs but missing from process document", expanded=False):
                        st.caption("These items exist in the source documents but were not captured in the generated process document. They do not need to be filled by the SME — consider re-running the document build to pick them up.")
                        for g in _doc_misses:
                            cat = g.get("category", "")
                            desc = g.get("description", "")
                            note = g.get("verification_note", "")
                            note_str = f" *(Source: {note})*" if note else ""
                            st.markdown(f"- `{cat}` — {desc}{note_str}")
            else:
                st.info("Run gap analysis to see ranked results.")

        with tab_categories:
            if gaps.gaps_by_category:
                categories = list(gaps.gaps_by_category.items())
                n_cols = min(3, len(categories))
                cols = st.columns(n_cols)
                for i, (cat, items) in enumerate(categories):
                    with cols[i % n_cols]:
                        st.subheader(cat)
                        for item in items:
                            _gap_item(item)
            else:
                st.info("No audience-specific gaps identified.")

        with tab_edge:
            if gaps.edge_cases:
                st.caption(f"{len(gaps.edge_cases)} edge cases identified")
                for item in gaps.edge_cases:
                    _gap_item(item)
            else:
                st.info("No edge cases identified.")

        with tab_resource:
            if gaps.resource_gaps:
                st.caption(f"{len(gaps.resource_gaps)} resource gaps identified")
                for item in gaps.resource_gaps:
                    _gap_item(item)
            else:
                st.info("No resource gaps identified.")

        st.divider()
        _aud_key = st.session_state.get("doc_audience", "new_employee")
        _aud_label = AUDIENCE_LABELS.get(_aud_key, _aud_key)
        st.download_button(
            "📥 Download Gap Analysis (Word)",
            data=_gap_analysis_to_word_bytes(
                gaps,
                st.session_state.process_document,
                st.session_state.current_session or "",
                audience_label=_aud_label,
            ),
            file_name="gap_analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


def _get_all_source_files() -> list:
    """Return all source files available for KB indexing from the current session state."""
    scanned = st.session_state.get("bulk_scanned")
    if scanned:
        return scanned.cobol + scanned.code + scanned.word + scanned.excel
    # Fall back to file lists derived from loaded clusters (session restore path)
    clusters = st.session_state.get("bulk_clusters")
    if clusters:
        seen: set[str] = set()
        files: list = []
        for cl in clusters:
            for f in cl.cobol_files + cl.doc_files + cl.shared_doc_files:
                key = str(f)
                if key not in seen:
                    seen.add(key)
                    files.append(f)
        if files:
            return files
    uploaded = st.session_state.get("uploaded_files", [])
    return [f for f in uploaded if isinstance(f, Path)]


def _render_chat_kb_banner(agent) -> None:
    """Show KB status banner and inline build/restart controls above the gap tracker."""
    _session = st.session_state.current_session
    if not _session:
        return

    _kb_dir = Path("./analysis_sessions") / _session / "knowledge_base"
    _kb = KnowledgeBase(_kb_dir)

    if not _kb.is_built():
        with st.expander(
            "💡 Knowledge Base not built — expand to build for richer responses",
            expanded=False,
        ):
            st.write(
                "Without the Knowledge Base the agent cannot link gaps to specific source "
                "passages, clarification responses won't include quoted excerpts, and document "
                "targeting relies on keyword overlap only.  Building typically takes 1–3 minutes."
            )
            all_files = _get_all_source_files()
            if not all_files:
                st.warning("No source files found in session state. Run the Bulk Load or upload documents first.")
                return

            if st.button("🏗️ Build Knowledge Base", key="gap_chat_build_kb", type="primary"):
                _prog = st.progress(0, text="Starting…")
                _status = st.empty()

                def _on_progress(msg, cur, tot):
                    _prog.progress(cur / tot if tot else 0, text=msg)
                    _status.write(msg)

                try:
                    _fcm: dict[str, str] = {}
                    _fcn: dict[str, str] = {}
                    for _cl in (st.session_state.get("bulk_clusters") or []):
                        for _f in list(getattr(_cl, "cobol_files", [])) + list(getattr(_cl, "doc_files", [])):
                            _fcm[_f.name] = _cl.cluster_id
                            _fcn[_f.name] = _cl.cluster_name

                    n = _kb.build(all_files, progress_callback=_on_progress,
                                  file_cluster_map=_fcm or None, file_cluster_names=_fcn or None)

                    _cs = st.session_state.get("bulk_cluster_summaries") or []
                    if _cs:
                        _status.write("Indexing cluster summaries…")
                        _kb.index_cluster_summaries(_cs)

                    _pd = st.session_state.get("process_document")
                    if _pd:
                        _status.write("Indexing process document…")
                        _kb.index_process_document(_pd)

                    _prog.progress(1.0, text="Done!")
                    _status.empty()
                    st.success(f"Knowledge Base built: {n} chunks from {len(all_files)} files.")
                    # Reset agent so it re-initialises with the fresh KB
                    st.session_state.chat_kb = _kb
                    st.session_state.chat_agent = None
                    st.session_state.chat_messages = []
                    st.session_state.document_updates = {}
                    st.rerun()
                except Exception as _e:
                    st.error(f"Build failed: {_e}")
    elif agent.knowledge_base is None:
        # KB was built (perhaps on the Knowledge Chat page) after this agent was initialised
        st.session_state.chat_kb = _kb  # keep session ref in sync
        st.info(
            "ℹ️ The Knowledge Base is now available. Restart the chat to enable source linking.",
        )
        if st.button("🔄 Restart Chat with Knowledge Base", key="gap_chat_restart_kb"):
            st.session_state.chat_agent = None
            st.session_state.chat_messages = []
            st.session_state.document_updates = {}
            st.rerun()
    else:
        # KB is built and the agent already has it — keep session ref current
        st.session_state.chat_kb = _kb


def render_chat_page():
    """Conversational gap-filling chat with the AI agent."""
    st.header("💬 Gap-Filling Chat")

    if not st.session_state.process_document:
        st.warning("Please generate the process document first.")
        return
    if not st.session_state.gap_analysis:
        st.warning("Please run gap analysis first.")
        return

    # Loaded session: messages exist but no live agent
    if st.session_state.chat_agent is None and st.session_state.chat_messages:
        st.info("Previous chat session loaded. Start a new session to continue the conversation.")
        for msg in st.session_state.chat_messages:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])
        if st.button("▶️ Start New Chat Session"):
            st.session_state.chat_messages = []
            st.session_state.document_updates = {}
            st.rerun()
        return

    # Initialize agent on first visit
    if st.session_state.chat_agent is None:
        with st.spinner("Initializing chat agent..."):
            # Load knowledge base for gap-to-source linking if it has been built
            _kb = None
            _session = st.session_state.current_session
            if _session:
                _kb_dir = Path("./analysis_sessions") / _session / "knowledge_base"
                _kb_candidate = KnowledgeBase(_kb_dir)
                if _kb_candidate.is_built():
                    _kb = _kb_candidate

            _audience = st.session_state.get("doc_audience", "new_employee")
            _sme_role = st.session_state.get("_sme_role_description", "").strip()
            _interview_note = AUDIENCE_INTERVIEW_NOTES.get(_audience, "")
            if _sme_role:
                _interview_note = (_interview_note + f"\n\nSpecific SME context: {_sme_role}").strip()
            agent = ConversationalAgent(
                analyses=st.session_state.analyses,
                process_document=st.session_state.process_document,
                gap_analysis=st.session_state.gap_analysis,
                knowledge_base=_kb,
                audience_note=_interview_note,
                response_mode=st.session_state.get("response_mode", "standard"),
                detail_level=st.session_state.get("detail_level", "standard"),
            )
            opening = agent.get_opening_message()
            st.session_state.chat_agent = agent
            st.session_state.chat_messages = [{"role": "assistant", "content": opening}]
            st.session_state.document_updates = {}
            st.session_state.chat_kb = _kb  # persist KB ref for update-after-doc-change
        st.rerun()

    agent: ConversationalAgent = st.session_state.chat_agent

    # ── KB status banner ─────────────────────────────────────────────────────
    _render_chat_kb_banner(agent)

    total_gaps = len(agent.gap_queue)
    resolved_gaps = sum(1 for g in agent.gap_queue if g.resolved)

    # ── Gap tracker (scrollable box, sits above the chat) ────────────────────
    if total_gaps > 0:
        st.subheader(f"Gaps  ({resolved_gaps}/{total_gaps} resolved)")

        # Map each gap to its most relevant cluster (word-overlap, no LLM call)
        gap_cluster_map = _map_gaps_to_clusters(
            [g.description for g in agent.gap_queue],
            st.session_state.analyses,
        )

        by_category: dict = defaultdict(list)
        for g in agent.gap_queue:
            by_category[g.category].append(g)

        import html as _html
        gap_rows = []
        for category, items in by_category.items():
            gap_rows.append(
                f"<div style='font-size:0.75rem;font-weight:600;color:#888;"
                f"text-transform:uppercase;margin-top:8px;padding:2px 0'>"
                f"{_html.escape(category.replace('_', ' '))}</div>"
            )
            for g in items:
                cluster = gap_cluster_map.get(g.description, "")
                if g.resolved:
                    icon = "✅"
                    row_style = "color:#4caf50;border-left:3px solid #4caf50;"
                    tip = (
                        f" title='{_html.escape(g.resolution[:120])}'"
                        if g.resolution else ""
                    )
                    gap_rows.append(
                        f"<div{tip} style='font-size:0.82rem;{row_style}"
                        f"padding:4px 0 4px 8px;margin:3px 0;line-height:1.4'>"
                        f"{icon}&nbsp;{_html.escape(g.description)}</div>"
                    )
                else:
                    icon = "⬜"
                    row_style = "color:#333;border-left:3px solid #ddd;"
                    rephrase_badge = (
                        "&nbsp;<span style='font-size:0.68rem;color:#888;font-style:italic'>(rephrased)</span>"
                        if g.rephrased else ""
                    )
                    gap_rows.append(
                        f"<div style='font-size:0.82rem;{row_style}"
                        f"padding:4px 0 2px 8px;margin:3px 0;line-height:1.4'>"
                        f"{icon}&nbsp;{_html.escape(g.description)}{rephrase_badge}</div>"
                    )
                    # Show cluster tag and source files under unresolved gaps
                    meta_parts = []
                    if cluster:
                        meta_parts.append(f"🏷️ {_html.escape(cluster)}")
                    if g.related_sources:
                        src_names = ", ".join(
                            _html.escape(s['file']) for s in g.related_sources[:3]
                        )
                        meta_parts.append(f"📎 {src_names}")
                    if meta_parts:
                        gap_rows.append(
                            f"<div style='font-size:0.72rem;color:#888;"
                            f"padding:0 0 4px 22px;line-height:1.3'>"
                            f"{'&nbsp;&nbsp;·&nbsp;&nbsp;'.join(meta_parts)}</div>"
                        )

        html_content = "\n".join(gap_rows)
        st.markdown(
            f"<div style='max-height:260px;overflow-y:auto;padding:6px 10px;"
            f"border:1px solid #e0e0e0;border-radius:6px;background:#fafafa;"
            f"margin-bottom:16px'>"
            f"{html_content}</div>",
            unsafe_allow_html=True,
        )

        # ── Gap management: rephrase or skip individual gaps ─────────────────
        _unresolved_gaps = [g for g in agent.gap_queue if not g.resolved]
        if _unresolved_gaps:
            with st.expander("Manage gaps — rephrase or skip"):
                st.caption(
                    "If the agent asks about a gap this SME wouldn't know, "
                    "select it here to rephrase it for their level or skip it entirely."
                )
                _mgmt_sel_col, _mgmt_btn_col = st.columns([3, 1])
                with _mgmt_sel_col:
                    _selected_idx = st.selectbox(
                        "Gap",
                        range(len(_unresolved_gaps)),
                        format_func=lambda i: (
                            f"{'↩ ' if _unresolved_gaps[i].rephrased else ''}"
                            f"[{_unresolved_gaps[i].category}] "
                            f"{_unresolved_gaps[i].description[:90]}"
                            f"{'…' if len(_unresolved_gaps[i].description) > 90 else ''}"
                        ),
                        label_visibility="collapsed",
                    )
                with _mgmt_btn_col:
                    _r_col, _s_col = st.columns(2)
                    with _r_col:
                        if st.button("🔄 Rephrase", help="Rewrite this gap to be answerable by the selected SME audience", use_container_width=True):
                            with st.spinner("Rephrasing…"):
                                agent.rephrase_gap(_unresolved_gaps[_selected_idx])
                            st.rerun()
                    with _s_col:
                        if st.button("✖ Skip", help="Mark this gap as not applicable to this SME", use_container_width=True):
                            agent.skip_gap(_unresolved_gaps[_selected_idx])
                            st.rerun()
                _sel_gap = _unresolved_gaps[_selected_idx]
                if _sel_gap.rephrased and _sel_gap.original_description:
                    st.caption(f"*Original: {_sel_gap.original_description[:120]}{'…' if len(_sel_gap.original_description) > 120 else ''}*")

    _render_chat_column(agent)


def _render_chat_column(agent) -> None:
    """Reset button, chat history, document updates sidebar, and chat input."""
    _gap_reset_col, _ = st.columns([1, 5])
    with _gap_reset_col:
        if st.button("🔄 Reset Chat"):
            st.session_state.chat_agent = None
            st.session_state.chat_messages = []
            st.session_state.document_updates = {}
            st.rerun()

    st.divider()

    # ── Chat history ──────────────────────────────────────────────────────────
    for msg in st.session_state.chat_messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    # ── Download section ──────────────────────────────────────────────────────
    _gap_msgs = st.session_state.chat_messages
    _gap_last = next((m for m in reversed(_gap_msgs) if m["role"] == "assistant"), None)
    if _gap_last:
        with st.expander("📥 Download"):
            st.caption("Latest response")
            _gd1, _gd2 = st.columns(2)
            with _gd1:
                st.download_button(
                    "📄 Word",
                    data=_chat_to_word_bytes([_gap_last], "Gap Chat Response"),
                    file_name="gap_chat_response.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="gap_dl_last_word",
                )
            with _gd2:
                st.download_button(
                    "📝 Text",
                    data=_chat_to_text([_gap_last], "Gap Chat Response"),
                    file_name="gap_chat_response.txt",
                    mime="text/plain",
                    key="gap_dl_last_txt",
                )
            st.caption("Full conversation")
            _gd3, _gd4 = st.columns(2)
            with _gd3:
                st.download_button(
                    "📄 Word",
                    data=_chat_to_word_bytes(_gap_msgs, "Gap Chat Conversation"),
                    file_name="gap_chat_conversation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="gap_dl_full_word",
                )
            with _gd4:
                st.download_button(
                    "📝 Text",
                    data=_chat_to_text(_gap_msgs, "Gap Chat Conversation"),
                    file_name="gap_chat_conversation.txt",
                    mime="text/plain",
                    key="gap_dl_full_txt",
                )

    # ── Document updates sidebar ──────────────────────────────────────────────
    if st.session_state.document_updates:
        with st.sidebar:
            st.divider()
            st.subheader("📄 Updated Documents")
            for doc_name, updates in st.session_state.document_updates.items():
                st.write(f"**{doc_name}** — {len(updates)} update(s)")
                for upd in updates:
                    label = f"v{upd.version}: {upd.gap_addressed[:35]}..."
                    with st.expander(label):
                        st.text_area(
                            "Content",
                            upd.content,
                            height=200,
                            key=f"upd_{doc_name}_{upd.version}",
                        )
                        st.download_button(
                            "📥 Download",
                            data=upd.content,
                            file_name=f"{doc_name}_v{upd.version}.md",
                            mime="text/markdown",
                            key=f"dl_{doc_name}_{upd.version}",
                        )

    # ── Input / completion ────────────────────────────────────────────────────
    if len(agent.remaining_gaps) == 0:
        st.success("✅ All gaps have been addressed! Download updated documents from the sidebar.")

        if st.session_state.document_updates:
            if st.button("💾 Save All Updates to Session"):
                session = _ensure_session()
                errors = []
                for doc_name, updates in st.session_state.document_updates.items():
                    for upd in updates:
                        try:
                            st.session_state.storage.save_document_update(
                                session,
                                upd.document_name,
                                upd.version,
                                upd.content,
                                upd.gap_addressed,
                            )
                        except Exception as e:
                            errors.append(f"{doc_name}: {e}")
                if errors:
                    st.error("Some saves failed: " + "; ".join(errors))
                else:
                    st.success(f"All document updates saved to session: **{session}**")
    else:
        if user_input := st.chat_input("Type your response..."):
            st.session_state.chat_messages.append({"role": "user", "content": user_input})

            with st.spinner("Thinking..."):
                response = agent.chat(user_input)

            st.session_state.chat_messages.append(
                {"role": "assistant", "content": response.message}
            )

            for update in response.document_updates:
                st.session_state.document_updates.setdefault(update.document_name, [])
                st.session_state.document_updates[update.document_name].append(update)
                try:
                    st.session_state.storage.save_document_update(
                        _ensure_session(),
                        update.document_name,
                        update.version,
                        update.content,
                        update.gap_addressed,
                    )
                except Exception:
                    pass

            # Refresh KB chunks for every updated document so future searches
            # reflect the newly incorporated information
            _chat_kb = st.session_state.get("chat_kb")
            if _chat_kb is not None and _chat_kb.is_built() and response.document_updates:
                for update in response.document_updates:
                    try:
                        _chat_kb.update_document_chunks(
                            source_file=f"[Updated Summary — {update.document_name}]",
                            new_content=update.content,
                        )
                    except Exception:
                        pass

            try:
                st.session_state.storage.save_chat_messages(
                    _ensure_session(),
                    st.session_state.chat_messages,
                )
            except Exception:
                pass

            st.rerun()



def _find_process_doc_excerpts(
    gap_desc: str,
    process_doc: "ProcessDocument",
    max_excerpts: int = 3,
) -> list[tuple[str, str]]:
    """Return up to max_excerpts (section_label, sentence) pairs from the process document
    that are most relevant to gap_desc, using keyword overlap scoring."""
    import re

    SECTIONS = [
        ("Overview", process_doc.overview),
        ("Integrated Processes", process_doc.integrated_processes),
        ("Dependencies", process_doc.dependencies),
        ("Data Flow", process_doc.data_flow),
        ("Decision Points", process_doc.decision_points),
        ("Systems & Components", process_doc.systems_and_components),
        ("Appendix", process_doc.appendix),
    ]

    stop_words = {
        "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
        "of", "with", "is", "are", "was", "were", "be", "been", "it", "this",
        "that", "as", "by", "from", "not", "no", "have", "has", "had", "will",
    }
    gap_words = {
        w.lower() for w in re.findall(r"\b\w+\b", gap_desc) if w.lower() not in stop_words and len(w) > 3
    }
    if not gap_words:
        return []

    scored: list[tuple[float, str, str]] = []
    for label, text in SECTIONS:
        if not text:
            continue
        sentences = re.split(r"(?<=[.!?])\s+", text)
        for sent in sentences:
            sent = sent.strip()
            if len(sent) < 20:
                continue
            sent_words = {w.lower() for w in re.findall(r"\b\w+\b", sent) if w.lower() not in stop_words}
            overlap = len(gap_words & sent_words) / max(len(gap_words), 1)
            if overlap > 0:
                scored.append((overlap, label, sent))

    scored.sort(key=lambda x: x[0], reverse=True)
    seen: set[str] = set()
    results: list[tuple[str, str]] = []
    for _, label, sent in scored:
        key = sent[:80]
        if key not in seen:
            seen.add(key)
            results.append((label, sent))
        if len(results) >= max_excerpts:
            break
    return results


def _gap_analysis_to_word_bytes(
    gap_analysis: "GapAnalysis",
    process_doc: "ProcessDocument",
    session_name: str,
    audience_label: str = "",
) -> bytes:
    """Build a Word document from gap analysis results with source passages."""
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_bg(cell, hex_color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _importance_label(score: int) -> tuple[str, str]:
        if score >= 8:
            return "Critical", "FF4444"
        elif score >= 6:
            return "High", "FF8C00"
        elif score >= 4:
            return "Medium", "FFD700"
        else:
            return "Low", "4CAF50"

    d = DocxDocument()
    d.add_heading("Gap Analysis Report", 0)

    meta = d.add_paragraph()
    meta.add_run(f"Session: ").bold = True
    meta.add_run(session_name or "—")
    meta.add_run(f"   |   Date: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d"))
    if audience_label:
        meta.add_run(f"   |   Audience: ").bold = True
        meta.add_run(audience_label)

    # ── Ranked Gaps ──────────────────────────────────────────────────────────
    _true_ranked = [g for g in gap_analysis.ranked_gaps if not g.get("is_doc_miss", False)]
    _miss_ranked = [g for g in gap_analysis.ranked_gaps if g.get("is_doc_miss", False)]

    d.add_heading("Ranked Gaps", level=1)
    if _true_ranked:
        for g in _true_ranked:
            score = g.get("importance", 5)
            cat = g.get("category", "")
            desc = g.get("description", "")
            level_label, hex_color = _importance_label(score)

            tbl = d.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr_row = tbl.rows[0]
            hdr_row.cells[0].text = f"{score}/10"
            hdr_row.cells[1].text = level_label
            hdr_row.cells[2].text = cat
            _set_cell_bg(hdr_row.cells[0], hex_color)
            _set_cell_bg(hdr_row.cells[1], hex_color)
            for cell in hdr_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(9)

            desc_para = d.add_paragraph(desc)
            desc_para.paragraph_format.left_indent = Pt(12)

            excerpts = _find_process_doc_excerpts(desc, process_doc)
            if excerpts:
                src_heading = d.add_paragraph()
                src_run = src_heading.add_run("Source passages:")
                src_run.italic = True
                src_run.font.size = Pt(9)
                src_heading.paragraph_format.left_indent = Pt(12)
                for section_label, passage in excerpts:
                    p = d.add_paragraph(style="List Bullet")
                    p.paragraph_format.left_indent = Pt(24)
                    lbl = p.add_run(f"[{section_label}] ")
                    lbl.bold = True
                    lbl.font.size = Pt(9)
                    txt = p.add_run(passage)
                    txt.font.size = Pt(9)

            d.add_paragraph()

    # ── Documentation Misses ──────────────────────────────────────────────────
    if _miss_ranked:
        d.add_heading("Documentation Misses", level=1)
        d.add_paragraph(
            "These items are covered in the source documents but were not captured in the "
            "generated process document. They do not need SME input — consider re-running "
            "the document build to pick them up."
        )
        for g in _miss_ranked:
            cat = g.get("category", "")
            desc = g.get("description", "")
            note = g.get("verification_note", "")
            p = d.add_paragraph(style="List Bullet")
            p.add_run(f"[{cat}] ").bold = True
            p.add_run(desc)
            if note:
                note_p = d.add_paragraph(f"Source: {note}")
                note_p.paragraph_format.left_indent = Pt(24)
                for r in note_p.runs:
                    r.italic = True
                    r.font.size = Pt(9)

    # ── Edge Cases ────────────────────────────────────────────────────────────
    if gap_analysis.edge_cases:
        d.add_heading("Edge Cases", level=1)
        for ec in gap_analysis.edge_cases:
            p = d.add_paragraph(style="List Bullet")
            p.add_run(ec)
            excerpts = _find_process_doc_excerpts(ec, process_doc, max_excerpts=2)
            if excerpts:
                for section_label, passage in excerpts:
                    sub = d.add_paragraph(style="List Bullet 2")
                    lbl = sub.add_run(f"[{section_label}] ")
                    lbl.bold = True
                    lbl.font.size = Pt(9)
                    txt = sub.add_run(passage)
                    txt.font.size = Pt(9)

    # ── Resource Gaps ─────────────────────────────────────────────────────────
    if gap_analysis.resource_gaps:
        d.add_heading("Resource Gaps", level=1)
        for rg in gap_analysis.resource_gaps:
            p = d.add_paragraph(style="List Bullet")
            p.add_run(rg)

    # ── Reference Index ───────────────────────────────────────────────────────
    # Collect every [ref:N] or bare [N] marker that appears in gap descriptions
    # and excerpt passages, then resolve them against the process doc citations.
    import re as _re
    _ref_scan_re = _re.compile(r'\[ref:(\d+)[^\]]*\]|\[(\d+)\]', _re.IGNORECASE)
    _cited_nums: set[str] = set()

    def _scan_refs(text: str) -> None:
        for _m in _ref_scan_re.finditer(text):
            _n = _m.group(1) or _m.group(2)
            if _n:
                _cited_nums.add(_n)

    for _g in gap_analysis.ranked_gaps:
        _scan_refs(_g.get("description", ""))
        for _, _passage in _find_process_doc_excerpts(_g.get("description", ""), process_doc):
            _scan_refs(_passage)
    for _ec in gap_analysis.edge_cases:
        _scan_refs(_ec)

    _resolved = {
        n: process_doc.citations[f"ref:{n}"]
        for n in _cited_nums
        if f"ref:{n}" in process_doc.citations
    } if process_doc and process_doc.citations else {}

    if _resolved:
        d.add_heading("References", level=1)
        d.add_paragraph(
            "The following source clusters are cited in gap descriptions and source passages above."
        )
        for _n, _info in sorted(_resolved.items(), key=lambda kv: int(kv[0])):
            _cluster_name = _info.get("cluster_name", f"ref:{_n}") if isinstance(_info, dict) else f"ref:{_n}"
            _files = _info.get("files", []) if isinstance(_info, dict) else []
            _rp = d.add_paragraph(style="List Bullet")
            _rl = _rp.add_run(f"[{_n}]  {_cluster_name}")
            _rl.bold = True
            if _files:
                _rp.add_run(" — " + ", ".join(_files))

    _add_footer_page_numbers(d)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _add_footer_page_numbers(d) -> None:
    """Add a right-aligned page number to the footer of every section."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in d.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        run._r.append(begin)
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        run._r.append(instr)
        end_el = OxmlElement("w:fldChar")
        end_el.set(qn("w:fldCharType"), "end")
        run._r.append(end_el)


def _add_content_to_docx(docx, content: str) -> None:
    """Convert LLM markdown-style output into Word paragraphs."""
    import re
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            docx.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            docx.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            docx.add_heading(stripped[2:].strip(), level=2)
        elif re.match(r"^[-*•]\s+", stripped):
            text = re.sub(r"^[-*•]\s+", "", stripped)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            docx.add_paragraph(text, style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", stripped):
            text = re.sub(r"^\d+[.)]\s+", "", stripped)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            docx.add_paragraph(text, style="List Number")
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            docx.add_paragraph(text)


@st.cache_data(show_spinner=False)
def _mermaid_to_png(mermaid_code: str) -> bytes | None:
    """Fetch a PNG render of Mermaid syntax from the mermaid.ink public API."""
    import base64
    import urllib.request
    clean = mermaid_code.strip()
    for fence in ("```mermaid", "```"):
        if clean.startswith(fence):
            clean = clean[len(fence):]
    if clean.endswith("```"):
        clean = clean[:-3]
    clean = clean.strip()
    try:
        encoded = base64.urlsafe_b64encode(clean.encode()).decode()
        url = f"https://mermaid.ink/img/{encoded}?bgColor=white"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            if resp.status == 200:
                return resp.read()
    except Exception:
        pass
    return None


def _strip_markdown(text: str) -> str:
    """Strip common markdown syntax for clean plain-text output."""
    import re
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)   # headers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)                  # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)                       # italic
    text = re.sub(r'`(.+?)`', r'\1', text)                         # inline code
    text = re.sub(r'^\s*[-*]\s+', '• ', text, flags=re.MULTILINE) # bullets
    return text


def _chat_to_text(messages: list[dict], title: str = "Chat Conversation") -> str:
    """Render a list of chat messages as plain text (markdown stripped)."""
    lines = [title, "=" * len(title), ""]
    for msg in messages:
        prefix = "You" if msg["role"] == "user" else "Assistant"
        lines.append(f"{prefix}:")
        lines.append(_strip_markdown(msg["content"]))
        lines.append("")
    return "\n".join(lines)


def _chat_to_word_bytes(messages: list[dict], title: str = "Chat Conversation") -> bytes:
    """Render a list of chat messages as a Word document, converting markdown to Word styles."""
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor

    d = DocxDocument()
    d.add_heading(title, 0)

    for msg in messages:
        is_user = msg["role"] == "user"
        label_para = d.add_paragraph()
        label_run = label_para.add_run("You:" if is_user else "Assistant:")
        label_run.bold = True
        label_run.font.size = Pt(11)
        if is_user:
            label_run.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

        _add_content_to_docx(d, msg["content"])
        d.add_paragraph()  # blank line between messages

    _add_footer_page_numbers(d)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _diagram_layout(png_bytes: bytes) -> "tuple[float, str | None, bool]":
    """Return (width_inches, paper_note, use_tabloid_landscape) for a rendered PNG.

    Tier 1 — portrait letter  : fits at 6.5" wide within 9" tall  → no note
    Tier 2 — portrait legal   : fits at 6.5" wide within 12.5" tall → suggest legal paper
    Tier 3 — tabloid landscape: everything else → landscape 11×17 section + paper note
    Falls back to (6.5, None, False) if the PNG header cannot be read.
    """
    try:
        import struct
        if len(png_bytes) < 24:
            return 6.5, None, False
        img_w, img_h = struct.unpack('>II', png_bytes[16:24])
        if img_w == 0:
            return 6.5, None, False
        aspect = img_h / img_w
    except Exception:
        return 6.5, None, False

    h_at_portrait = 6.5 * aspect
    if h_at_portrait <= 9.0:
        return 6.5, None, False
    elif h_at_portrait <= 12.5:
        return 6.5, 'For best results, print on legal paper (8.5" × 14").', False
    else:
        # Landscape tabloid: 16" × 9.5" usable (17" × 11" with 0.5" margins)
        pic_w = min(16.0, 9.5 / aspect)
        return pic_w, 'For best results, print on 11" × 17" paper (landscape).', True


def _insert_page_section(d, w_twips: str, h_twips: str, margin: str, orient: "str | None" = None) -> None:
    """Append a section-break paragraph that ends the current section with the given page geometry."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = d.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), w_twips)
    pgSz.set(qn('w:h'), h_twips)
    if orient:
        pgSz.set(qn('w:orient'), orient)
    sectPr.append(pgSz)
    pgMar = OxmlElement('w:pgMar')
    for attr in ('w:top', 'w:right', 'w:bottom', 'w:left'):
        pgMar.set(qn(attr), margin)
    sectPr.append(pgMar)
    pPr.append(sectPr)


def _generate_word_doc(doc) -> bytes:
    """Build a Word document from a ProcessDocument and return raw bytes."""
    import io
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt

    d = DocxDocument()
    d.add_heading("Process Document", 0)

    sections = [
        ("Overview", doc.overview),
        ("Integrated Processes", doc.integrated_processes),
        ("Dependencies", doc.dependencies),
        ("Data Flow", doc.data_flow),
        ("Decision Points", doc.decision_points),
        ("Systems & Components", doc.systems_and_components),
    ]
    if doc.appendix:
        sections.append(("Appendix", doc.appendix))

    for title, content in sections:
        d.add_heading(title, level=1)
        _add_content_to_docx(d, content)

    if doc.process_flow_diagram:
        d.add_heading("Process Flow Diagram", level=1)
        png = _mermaid_to_png(doc.process_flow_diagram)
        if png:
            pic_w_in, paper_note, use_landscape = _diagram_layout(png)
            if use_landscape:
                _insert_page_section(d, '12240', '15840', '1440')
            d.add_picture(io.BytesIO(png), width=Inches(pic_w_in))
            if paper_note:
                note_p = d.add_paragraph(paper_note)
                for r in note_p.runs:
                    r.italic = True
                    r.font.size = Pt(9)
            if use_landscape:
                _insert_page_section(d, '24480', '15840', '720', orient='landscape')
        else:
            d.add_paragraph(
                "Diagram could not be rendered. "
                "Paste the Mermaid source below at https://mermaid.live to view it."
            )
        # Always embed Mermaid source so the diagram is reproducible
        src_label = d.add_paragraph("Mermaid source (paste at https://mermaid.live):")
        for r in src_label.runs:
            r.italic = True
            r.font.size = Pt(9)
        code_para = d.add_paragraph(doc.process_flow_diagram)
        for run in code_para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    if doc.citations:
        import re as _re
        d.add_heading("Citations", level=1)
        d.add_paragraph(
            "The following source clusters were referenced during document generation. "
            "Inline [N] markers in the document body trace to the cluster listed here."
        )
        for ref_key, info in sorted(doc.citations.items(), key=lambda kv: kv[0]):
            cluster_name = info.get("cluster_name", ref_key) if isinstance(info, dict) else ref_key
            files = info.get("files", []) if isinstance(info, dict) else []
            # ref_key is "ref:N" — display as [N] to match inline markers
            _m = _re.match(r'ref:(\d+)', ref_key)
            label_num = _m.group(1) if _m else ref_key
            p = d.add_paragraph(style="List Bullet")
            label = p.add_run(f"[{label_num}]  {cluster_name}")
            label.bold = True
            if files:
                p.add_run(" — " + ", ".join(files))

    _add_footer_page_numbers(d)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _render_mermaid(diagram_code: str, height: int = 500) -> None:
    """Render a Mermaid diagram using the Mermaid.js CDN via st.components."""
    # Strip markdown code fences if the LLM wrapped the output
    clean = diagram_code.strip()
    for fence in ("```mermaid", "```"):
        if clean.startswith(fence):
            clean = clean[len(fence):]
    if clean.endswith("```"):
        clean = clean[:-3]
    clean = clean.strip()

    html = f"""
<!DOCTYPE html>
<html>
<head>
  <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
  <script>mermaid.initialize({{startOnLoad:true, theme:'default', securityLevel:'loose'}});</script>
  <style>
    body {{ margin: 0; padding: 8px; background: #fff; }}
    .mermaid {{ width: 100%; overflow: auto; }}
  </style>
</head>
<body>
  <div class="mermaid">{clean}</div>
</body>
</html>"""
    st.components.v1.html(html, height=height, scrolling=True)


def render_knowledge_chat_page() -> None:
    """RAG-powered Q&A over the original source documents."""
    st.header("🔎 Knowledge Chat")
    st.write(
        "Ask questions about the process — answers are pulled directly from the "
        "original source files (Word, Excel, code). "
        "This is separate from the document-generation pipeline."
    )

    # Determine which files are available to index
    scanned = st.session_state.get("bulk_scanned")
    all_files = []
    if scanned:
        all_files = scanned.cobol + scanned.code + scanned.word + scanned.excel
    elif st.session_state.uploaded_files:
        all_files = [
            f for f in st.session_state.uploaded_files if isinstance(f, Path)
        ]

    if not all_files:
        st.warning(
            "No documents loaded yet. Go to **Bulk Load** or **Upload Documents** first, "
            "then come back here to build the knowledge base."
        )
        return

    # ── Knowledge base status ─────────────────────────────────────────
    session = st.session_state.current_session
    kb_dir = (
        Path("./analysis_sessions") / session / "knowledge_base"
        if session
        else Path("./knowledge_base_temp")
    )
    kb = KnowledgeBase(kb_dir)

    if not kb.is_built():
        st.info(
            f"The knowledge base has not been built yet. "
            f"{len(all_files)} files are ready to index."
        )
        if st.button("🏗️ Build Knowledge Base", type="primary"):
            progress_bar = st.progress(0, text="Starting…")
            status_text = st.empty()

            def on_kb_progress(msg, current, total):
                frac = current / total if total else 0
                progress_bar.progress(frac, text=msg)
                status_text.write(msg)

            try:
                # Build file→cluster maps so each chunk is stamped with its cluster
                _file_cluster_map: dict[str, str] = {}
                _file_cluster_names: dict[str, str] = {}
                _clusters = st.session_state.get("bulk_clusters") or []
                for _cl in _clusters:
                    for _f in list(getattr(_cl, "cobol_files", [])) + list(getattr(_cl, "doc_files", [])):
                        _file_cluster_map[_f.name] = _cl.cluster_id
                        _file_cluster_names[_f.name] = _cl.cluster_name

                n_chunks = kb.build(
                    all_files,
                    progress_callback=on_kb_progress,
                    file_cluster_map=_file_cluster_map or None,
                    file_cluster_names=_file_cluster_names or None,
                )

                # Index cluster summaries as searchable plain-English descriptions
                _cluster_summaries = st.session_state.get("bulk_cluster_summaries") or []
                summary_chunks = 0
                if _cluster_summaries:
                    status_text.write("Indexing cluster summaries…")
                    summary_chunks = kb.index_cluster_summaries(_cluster_summaries)

                # Also index the generated process document as tier-2
                proc_doc = st.session_state.get("process_document")
                extra_chunks = 0
                if proc_doc:
                    status_text.write("Indexing generated process document…")
                    extra_chunks = kb.index_process_document(proc_doc)

                progress_bar.progress(1.0, text="Done!")
                status_text.empty()
                extra_parts = []
                if summary_chunks:
                    extra_parts.append(f"{summary_chunks} cluster-summary chunks")
                if extra_chunks:
                    extra_parts.append(f"{extra_chunks} process-doc chunks")
                extra_msg = (" + " + ", ".join(extra_parts)) if extra_parts else ""
                st.success(
                    f"Knowledge base built: {n_chunks} chunks from {len(all_files)} files{extra_msg}."
                )
                # Initialise agent immediately
                _audience = st.session_state.get("doc_audience", "new_employee")
                st.session_state.rag_agent = RAGAgent(
                    kb,
                    audience_note=AUDIENCE_NOTES.get(_audience, ""),
                    response_mode=st.session_state.get("response_mode", "standard"),
                    detail_level=st.session_state.get("detail_level", "standard"),
                )
                st.session_state.rag_messages = []
                st.rerun()
            except Exception as e:
                st.error(f"Failed to build knowledge base: {e}")
        return

    # ── Initialise agent (or reinitialise if audience changed) ───────
    _audience = st.session_state.get("doc_audience", "new_employee")
    if _audience == "custom":
        _current_audience_note = (
            f"IMPORTANT — Audience: {st.session_state.get('custom_audience_note', '').strip()}"
            if st.session_state.get("custom_audience_note", "").strip() else ""
        )
    else:
        _current_audience_note = AUDIENCE_NOTES.get(_audience, "")

    _existing_agent = st.session_state.rag_agent
    _current_mode = st.session_state.get("response_mode", "standard")
    if _existing_agent is None or getattr(_existing_agent, "_audience_note", None) != _current_audience_note:
        proc_doc = st.session_state.get("process_document")
        if _existing_agent is None and proc_doc:
            kb.index_process_document(proc_doc)
        st.session_state.rag_agent = RAGAgent(
            kb,
            audience_note=_current_audience_note,
            response_mode=_current_mode,
            detail_level=st.session_state.get("detail_level", "standard"),
        )
        if _existing_agent is not None:
            # Preserve conversation history when only the audience changed
            st.session_state.rag_agent._history = _existing_agent._history
        stats = kb.get_stats()
        if _existing_agent is None:
            st.success(
                f"Knowledge base loaded: {stats['chunks']} chunks from {stats['files']} files."
            )

    agent: RAGAgent = st.session_state.rag_agent
    stats = kb.get_stats()

    col1, col2, col3 = st.columns(3)
    col1.metric("Source files", stats["files"])
    col2.metric("Index chunks", stats["chunks"])
    col3.metric("Messages", sum(1 for m in st.session_state.rag_messages if m["role"] == "user"))

    _mode_col, _detail_col, _clear_col, _rebuild_col = st.columns([2, 2, 1, 1])
    with _mode_col:
        st.radio(
            "Response mode",
            options=list(RESPONSE_MODE_LABELS.keys()),
            format_func=lambda k: RESPONSE_MODE_LABELS[k],
            index=list(RESPONSE_MODE_LABELS.keys()).index(_current_mode),
            horizontal=True,
            key="response_mode",
            help=(
                "Quick Summary — brief overview only  |  "
                "Standard — balanced detail  |  "
                "Step-by-Step Guide — numbered actions with exact screen/field details"
            ),
        )
    with _detail_col:
        _current_detail = st.session_state.get("detail_level", "standard")
        st.radio(
            "Detail level",
            options=list(DETAIL_LEVEL_LABELS.keys()),
            format_func=lambda k: DETAIL_LEVEL_LABELS[k],
            index=list(DETAIL_LEVEL_LABELS.keys()).index(_current_detail),
            horizontal=True,
            key="detail_level",
            help=(
                "Overview — major phases only  |  "
                "Standard — named screens and key fields  |  "
                "In-Depth — every field, exact values, system responses"
            ),
        )
    # Sync both controls to the live agent without reinitialising
    agent._response_mode = st.session_state.get("response_mode", "standard")
    agent._detail_level = st.session_state.get("detail_level", "standard")
    with _clear_col:
        if st.button("🗑️ Clear Conversation"):
            agent.reset()
            st.session_state.rag_messages = []
            st.rerun()
    with _rebuild_col:
        if st.button("🔄 Rebuild KB"):
            import shutil
            if kb_dir.exists():
                shutil.rmtree(kb_dir)
            st.session_state.rag_agent = None
            st.session_state.rag_messages = []
            st.rerun()

    st.divider()

    # ── Chat history ──────────────────────────────────────────────────
    for msg in st.session_state.rag_messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg.get("enriched_query") and msg["enriched_query"] != msg.get("raw_query"):
                st.caption(f"🔍 Searched for: *{msg['enriched_query']}*")
            if msg.get("sources"):
                with st.expander("📄 Source excerpts used"):
                    for src in msg["sources"]:
                        score_pct = int(src["score"] * 100)
                        st.markdown(
                            f"**{src['source_file']}** — relevance {score_pct}%\n\n"
                            f"> {src['text'][:1000]}{'…' if len(src['text']) > 1000 else ''}"
                        )

    # ── Download section ───────────────────────────────────────────────
    _rag_msgs = st.session_state.rag_messages
    _last_assistant = next(
        (m for m in reversed(_rag_msgs) if m["role"] == "assistant"), None
    )
    if _last_assistant:
        with st.expander("📥 Download"):
            st.caption("Latest response")
            _dl1, _dl2 = st.columns(2)
            with _dl1:
                st.download_button(
                    "📄 Word",
                    data=_chat_to_word_bytes([_last_assistant], "Knowledge Chat Response"),
                    file_name="knowledge_chat_response.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="rag_dl_last_word",
                )
            with _dl2:
                st.download_button(
                    "📝 Text",
                    data=_chat_to_text([_last_assistant], "Knowledge Chat Response"),
                    file_name="knowledge_chat_response.txt",
                    mime="text/plain",
                    key="rag_dl_last_txt",
                )
            st.caption("Full conversation")
            _dl3, _dl4 = st.columns(2)
            with _dl3:
                st.download_button(
                    "📄 Word",
                    data=_chat_to_word_bytes(_rag_msgs, "Knowledge Chat Conversation"),
                    file_name="knowledge_chat_conversation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="rag_dl_full_word",
                )
            with _dl4:
                st.download_button(
                    "📝 Text",
                    data=_chat_to_text(_rag_msgs, "Knowledge Chat Conversation"),
                    file_name="knowledge_chat_conversation.txt",
                    mime="text/plain",
                    key="rag_dl_full_txt",
                )

    # ── Input ──────────────────────────────────────────────────────────
    if user_input := st.chat_input("Ask anything about the process or documents…"):
        st.session_state.rag_messages.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.write(user_input)

        with st.chat_message("assistant"):
            with st.spinner("Enriching query and searching knowledge base…"):
                try:
                    answer, sources, enriched_query = agent.chat(user_input)
                except Exception as e:
                    answer = f"Error generating answer: {e}"
                    sources = []
                    enriched_query = user_input
            st.markdown(answer)
            if enriched_query and enriched_query != user_input:
                st.caption(f"🔍 Searched for: *{enriched_query}*")
            if sources:
                with st.expander("📄 Source excerpts used"):
                    for src in sources:
                        score_pct = int(src["score"] * 100)
                        st.markdown(
                            f"**{src['source_file']}** — relevance {score_pct}%\n\n"
                            f"> {src['text'][:300]}{'…' if len(src['text']) > 300 else ''}"
                        )

        st.session_state.rag_messages.append({
            "role": "assistant",
            "content": answer,
            "sources": sources,
            "enriched_query": enriched_query,
            "raw_query": user_input,
        })
        st.rerun()


def main():
    """Main application entry point."""
    st.set_page_config(
        page_title="Document Analysis System",
        page_icon="📚",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Apply custom styling
    st.markdown("""
        <style>
        .main {
            padding: 0rem 0rem;
        }
        </style>
    """, unsafe_allow_html=True)

    initialize_session()

    page = render_sidebar()

    # Main content
    if page == "Group Documents":
        render_group_documents_page()
    elif page == "Dependency Graph":
        render_dependency_graph_page()
    elif page == "Review Process Document":
        render_process_document_page()
    elif page == "Gap Analysis":
        render_gap_analysis_page()
    elif page == "Gap-Filling Chat":
        render_chat_page()
    elif page == "Knowledge Chat":
        render_knowledge_chat_page()

    # Footer
    st.divider()
    _, col2, _ = st.columns(3)
    with col2:
        st.caption("Document Analysis System v2.0 - With Session Management & Document Refinement")


if __name__ == "__main__":
    main()
